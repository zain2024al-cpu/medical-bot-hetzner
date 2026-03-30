# =============================
# bot/handlers/admin/admin_reports.py
# 🖨️ طباعة وفلترة التقارير مع الرسوم البيانية والإحصائيات ✅
# =============================
from telegram import Update, InlineKeyboardMarkup, InlineKeyboardButton
from telegram.ext import (
    ContextTypes,
    ConversationHandler,
    MessageHandler,
    CallbackQueryHandler,
    filters,
)
from datetime import datetime
from db.session import SessionLocal
from db.models import Report, Patient, Hospital, Department, Doctor
from services.pdf_generator import generate_pdf_report, generate_pdf_reports
from bot.shared_auth import is_admin
from bot.decorators import admin_handler
import os
import io
import asyncio
import base64
import logging
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from collections import Counter
import sys

logger = logging.getLogger(__name__)

# محاولة استيراد مكتبات العربية (اختيارية)
try:
    from arabic_reshaper import reshape
    from bidi.algorithm import get_display
    ARABIC_SUPPORT = True
except ImportError:
    ARABIC_SUPPORT = False

# ✅ حالات الفلترة المتقدمة (متوافقة مع admin_printing)
SELECT_FILTER, ENTER_NAME, SELECT_DEPARTMENT_OPTION, ENTER_DEPARTMENT, SELECT_YEAR, SELECT_MONTH, CONFIRM_EXPORT, PRINT_SELECT_PATIENT, SELECT_ACTION_TYPE = range(720, 729)


# ================================================
# ✅ معالجة اختيار المريض من Inline Query للطباعة
# ================================================

async def handle_print_patient_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """معالجة أمر /print_patient من Inline Query"""
    
    if not update.message or not update.message.text:
        return
    
    user = update.effective_user
    if not is_admin(user.id):
        return
    
    # استخراج اسم المريض من الأمر
    command_parts = update.message.text.split(maxsplit=1)
    if len(command_parts) < 2:
        await update.message.reply_text("❌ خطأ في اختيار المريض")
        return
    
    patient_name = command_parts[1].strip()
    
    # حفظ اسم المريض في context
    context.user_data["patient_name"] = patient_name
    context.user_data["filter_type"] = "patient"
    context.user_data["filter_value"] = patient_name
    
    # البحث عن المريض
    with SessionLocal() as session:
        patient = session.query(Patient).filter(
            Patient.full_name == patient_name
        ).first()
        
        if not patient:
            await update.message.reply_text(
                f"❌ لم يتم العثور على المريض: {patient_name}"
            )
            return
        
        # حساب عدد التقارير
        reports_count = session.query(Report).filter(
            Report.patient_id == patient.id
        ).count()
        
        # عرض خيارات الطباعة
        keyboard = InlineKeyboardMarkup([
            [InlineKeyboardButton("🖨️ طباعة كل التقارير", callback_data=f"print_patient_all")],
            [InlineKeyboardButton("📅 اختيار فترة محددة", callback_data=f"print_patient_period")],
            [InlineKeyboardButton("🔙 رجوع", callback_data="back:filter")],
            [InlineKeyboardButton("❌ إلغاء", callback_data="abort")]
        ])
        
        await update.message.reply_text(
            f"✅ **تم اختيار المريض**\n\n"
            f"👤 الاسم: {patient_name}\n"
            f"📊 عدد التقارير: {reports_count}\n\n"
            f"اختر خيار الطباعة:",
            reply_markup=keyboard,
            parse_mode="Markdown"
        )


async def handle_print_patient_options(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """معالجة خيارات طباعة تقارير المريض"""
    query = update.callback_query
    await query.answer()
    
    action = query.data
    patient_name = context.user_data.get("patient_name")
    
    if action == "print_patient_all":
        # طباعة كل التقارير مباشرة
        context.user_data["year_value"] = None
        context.user_data["month_value"] = None
        
        await query.edit_message_text(
            f"🖨️ **طباعة كل تقارير المريض**\n\n"
            f"👤 {patient_name}\n\n"
            f"✅ هل تريد المتابعة؟",
            reply_markup=_confirm_kb(),
            parse_mode="Markdown"
        )
        return CONFIRM_EXPORT
    
    elif action == "print_patient_period":
        # اختيار فترة محددة
        await query.edit_message_text(
            f"👤 **المريض:** {patient_name}\n\n"
            f"📅 اختر السنة:",
            reply_markup=_years_kb()
        )
        return SELECT_YEAR


async def handle_back_to_filter(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """معالجة زر الرجوع إلى قائمة الفلترة"""
    query = update.callback_query
    await query.answer()
    
    # مسح البيانات المؤقتة
    context.user_data.pop("mode", None)
    context.user_data.pop("patient_name", None)
    
    await query.edit_message_text("🖨️ اختر نوع الفلترة:", reply_markup=_filters_kb())
    return SELECT_FILTER


# ================================================
# دوال مساعدة للرسوم البيانية
# ================================================

def format_arabic_text(text):
    """تنسيق النص العربي للعرض الصحيح في الرسوم البيانية"""
    if not text:
        return text
    
    # محاولة استخدام المكتبات العربية
    if ARABIC_SUPPORT:
        try:
            # إعادة تشكيل النص العربي
            reshaped_text = reshape(text)
            # عكس اتجاه النص للعرض الصحيح
            bidi_text = get_display(reshaped_text)
            return bidi_text
        except Exception as e:
            print(f"خطأ في تنسيق النص العربي: {e}")
            return text
    
    # إذا لم تتوفر المكتبات، نرجع النص كما هو
    return text

def generate_statistics(reports, session):
    """إنشاء إحصائيات شاملة"""
    if not reports:
        return {}
    
    stats = {
        'total_reports': len(reports),
        'unique_patients': len(set(r.patient_id for r in reports if r.patient_id)),
        'unique_hospitals': len(set(r.hospital_id for r in reports if r.hospital_id)),
    }
    
    # التقسيم حسب المستشفى
    hospitals = {}
    for r in reports:
        if r.hospital_id:
            h = session.get(Hospital, r.hospital_id)
            hospital_name = h.name if h else 'غير محدد'
            hospitals[hospital_name] = hospitals.get(hospital_name, 0) + 1
    stats['by_hospital'] = hospitals
    
    # التقسيم حسب القسم
    departments = {}
    for r in reports:
        if r.department_id:
            d = session.get(Department, r.department_id)
            dept_name = d.name if d else 'غير محدد'
            departments[dept_name] = departments.get(dept_name, 0) + 1
    stats['by_department'] = departments
    
    # التقسيم حسب الإجراء الطبي
    actions = {}
    for r in reports:
        if r.medical_action:
            action = r.medical_action
            actions[action] = actions.get(action, 0) + 1
    stats['by_action'] = actions
    
    return stats

def create_charts(stats, filter_type=None):
    """
    إنشاء رسوم بيانية وتحويلها لـ base64
    filter_type: نوع الفلتر ('all', 'patient', 'hospital', 'date')
    """
    charts = {}
    
    if not stats:
        return charts
    
    # إعداد matplotlib للعربية
    plt.rcParams['font.family'] = ['Arial', 'DejaVu Sans', 'sans-serif']
    plt.rcParams['axes.unicode_minus'] = False
    plt.rcParams['font.size'] = 11
    
    # 1. رسم دائري: التوزيع حسب المستشفيات
    # يظهر فقط إذا:
    # - هناك أكثر من مستشفى واحد
    # - ونوع الفلتر ليس "hospital" (لأنه مستشفى محدد)
    if (stats.get('by_hospital') and 
        len(stats.get('by_hospital', {})) > 1 and 
        filter_type != 'hospital'):
        try:
            fig, ax = plt.subplots(figsize=(10, 8))
            
            hospitals = stats['by_hospital']
            labels = [format_arabic_text(label) for label in hospitals.keys()]
            values = list(hospitals.values())
            
            colors = ['#FF6384', '#36A2EB', '#FFCE56', '#4BC0C0', '#9966FF', '#FF9F40']
            
            # إضافة الأرقام والنسب
            def make_autopct(values):
                def my_autopct(pct):
                    total = sum(values)
                    val = int(round(pct*total/100.0))
                    return f'{val}\n({pct:.1f}%)'
                return my_autopct
            
            ax.pie(values, labels=labels, autopct=make_autopct(values), startangle=90, 
                   colors=colors[:len(labels)], textprops={'fontsize': 10, 'weight': 'bold'})
            ax.set_title(format_arabic_text('التوزيع حسب المستشفى'), 
                        pad=20, fontsize=16, fontweight='bold')
            
            # تحويل لـ base64
            buf = io.BytesIO()
            plt.tight_layout()
            plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
            buf.seek(0)
            charts['hospitals_pie'] = base64.b64encode(buf.read()).decode()
            plt.close()
        except Exception as e:
            print(f"خطأ في رسم المستشفيات: {e}")
    
    # 2. رسم بالأعمدة: التوزيع حسب الأقسام
    # يظهر فقط إذا هناك أكثر من قسم واحد
    if stats.get('by_department') and len(stats.get('by_department', {})) > 1:
        try:
            fig, ax = plt.subplots(figsize=(12, 6))
            
            departments = stats['by_department']
            # ترتيب حسب العدد
            sorted_depts = sorted(departments.items(), key=lambda x: x[1], reverse=True)[:10]
            
            labels = [format_arabic_text(label) for label, _ in sorted_depts]
            values = [value for _, value in sorted_depts]
            
            colors = ['#667eea', '#764ba2', '#f093fb', '#4facfe', '#00f2fe']
            bars = ax.barh(range(len(labels)), values, color=colors[0], height=0.6)
            
            # إضافة الأرقام على الأعمدة
            for i, (bar, value) in enumerate(zip(bars, values)):
                ax.text(value + max(values)*0.02, i, f'{value}', 
                       va='center', fontsize=11, fontweight='bold', color='#2c3e50')
            
            ax.set_yticks(range(len(labels)))
            ax.set_yticklabels(labels, fontsize=11)
            ax.set_xlabel(format_arabic_text('عدد التقارير'), fontsize=13, fontweight='bold')
            ax.set_title(format_arabic_text('التوزيع حسب الأقسام (أعلى 10)'), 
                        pad=20, fontsize=16, fontweight='bold')
            ax.grid(axis='x', alpha=0.3)
            ax.invert_yaxis()  # عكس المحور لعرض الأعلى في الأعلى
            
            # تحويل لـ base64
            buf = io.BytesIO()
            plt.tight_layout()
            plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
            buf.seek(0)
            charts['departments_bar'] = base64.b64encode(buf.read()).decode()
            plt.close()
        except Exception as e:
            print(f"خطأ في رسم الأقسام: {e}")
    
    # 3. رسم بالأعمدة: التوزيع حسب الإجراءات
    # يظهر فقط إذا هناك أكثر من إجراء واحد
    if stats.get('by_action') and len(stats.get('by_action', {})) > 1:
        try:
            fig, ax = plt.subplots(figsize=(12, 6))
            
            actions = stats['by_action']
            # ترتيب حسب العدد
            sorted_actions = sorted(actions.items(), key=lambda x: x[1], reverse=True)[:10]
            
            labels = [format_arabic_text(label[:30] + '...' if len(label) > 30 else label) 
                     for label, _ in sorted_actions]
            values = [value for _, value in sorted_actions]
            
            colors = ['#27ae60', '#229954', '#1e8449', '#196f3d', '#145a32']
            bars = ax.bar(range(len(labels)), values, color=colors[0], width=0.6)
            
            # إضافة الأرقام على الأعمدة
            for i, (bar, value) in enumerate(zip(bars, values)):
                ax.text(i, value + max(values)*0.02, f'{value}', 
                       ha='center', va='bottom', fontsize=11, fontweight='bold', color='#2c3e50')
            
            ax.set_xticks(range(len(labels)))
            ax.set_xticklabels(labels, rotation=45, ha='right', fontsize=10)
            ax.set_ylabel(format_arabic_text('عدد التقارير'), fontsize=13, fontweight='bold')
            ax.set_title(format_arabic_text('التوزيع حسب الإجراءات الطبية (أعلى 10)'), 
                        pad=20, fontsize=16, fontweight='bold')
            ax.grid(axis='y', alpha=0.3)
            
            # تحويل لـ base64
            buf = io.BytesIO()
            plt.tight_layout()
            plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
            buf.seek(0)
            charts['actions_bar'] = base64.b64encode(buf.read()).decode()
            plt.close()
        except Exception as e:
            print(f"خطأ في رسم الإجراءات: {e}")
    
    return charts

def _cancel_inline():
    return InlineKeyboardMarkup([[InlineKeyboardButton("❌ إلغاء المحادثة", callback_data="abort")]])


def _build_patient_list_keyboard(patients_list, page=0, items_per_page=10):
    """بناء قائمة أزرار المرضى مع صفحات"""
    total = len(patients_list)
    total_pages = max(1, (total + items_per_page - 1) // items_per_page)
    page = max(0, min(page, total_pages - 1))
    
    start_idx = page * items_per_page
    end_idx = min(start_idx + items_per_page, total)
    
    keyboard = []
    
    # عرض المرضى في الصفحة الحالية (اسمين في كل صف)
    patients_page = patients_list[start_idx:end_idx]
    for i in range(0, len(patients_page), 2):
        row = []
        for j in range(2):
            if i + j >= len(patients_page):
                break
            patient_id, patient_name = patients_page[i + j]
            # اختصار الاسم إذا كان طويلاً
            display_name = patient_name[:25] + "..." if len(patient_name) > 25 else patient_name
            row.append(InlineKeyboardButton(
                f"👤 {display_name}",
                callback_data=f"print_patient:{patient_id}"
            ))
        keyboard.append(row)
    
    # أزرار التنقل بين الصفحات
    nav_buttons = []
    if page > 0:
        nav_buttons.append(InlineKeyboardButton("⬅️ السابق", callback_data=f"patient_page:{page-1}"))
    nav_buttons.append(InlineKeyboardButton(f"📄 {page+1}/{total_pages}", callback_data="noop"))
    if page < total_pages - 1:
        nav_buttons.append(InlineKeyboardButton("➡️ التالي", callback_data=f"patient_page:{page+1}"))
    
    if nav_buttons:
        keyboard.append(nav_buttons)
    
    # أزرار التحكم
    keyboard.append([InlineKeyboardButton("🔙 رجوع", callback_data="back:filter")])
    keyboard.append([InlineKeyboardButton("❌ إلغاء", callback_data="abort")])
    
    text = (
        f"👤 **اختر المريض من القائمة:**\n\n"
        f"📊 إجمالي المرضى: {total}\n"
        f"📄 الصفحة: {page + 1} من {total_pages}"
    )
    
    return text, InlineKeyboardMarkup(keyboard)


def _filters_kb():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("👤 طباعة باسم مريض (إجراء + مدة)", callback_data="filter:patient")],
        [InlineKeyboardButton("🏥 طباعة باسم مستشفى", callback_data="filter:hospital")],
        [InlineKeyboardButton("🏢 طباعة حسب القسم", callback_data="filter:department")],
        [InlineKeyboardButton("📅 طباعة حسب التاريخ", callback_data="filter:date")],
        [InlineKeyboardButton("📄 طباعة كل التقارير", callback_data="filter:all")],
        [InlineKeyboardButton("❌ إلغاء", callback_data="abort")],
    ])

def _years_kb(show_back=True):
    """لوحة اختيار السنوات - من 2025 فصاعداً"""
    current_year = datetime.now().year
    start_year = 2025  # ✅ البداية من 2025
    years = []
    
    # من السنة الحالية إلى 2025
    for year in range(current_year, start_year - 1, -1):
        years.append([InlineKeyboardButton(f"📅 {year}", callback_data=f"year:{year}")])
    
    years.append([InlineKeyboardButton("📄 الكل (من 2025)", callback_data="year:all")])
    if show_back:
        years.append([InlineKeyboardButton("🔙 رجوع", callback_data="back:filter")])
    years.append([InlineKeyboardButton("❌ إلغاء", callback_data="abort")])
    return InlineKeyboardMarkup(years)

def _months_kb():
    """لوحة اختيار الشهور"""
    months = [
        [InlineKeyboardButton("يناير", callback_data="month:1"), InlineKeyboardButton("فبراير", callback_data="month:2")],
        [InlineKeyboardButton("مارس", callback_data="month:3"), InlineKeyboardButton("أبريل", callback_data="month:4")],
        [InlineKeyboardButton("مايو", callback_data="month:5"), InlineKeyboardButton("يونيو", callback_data="month:6")],
        [InlineKeyboardButton("يوليو", callback_data="month:7"), InlineKeyboardButton("أغسطس", callback_data="month:8")],
        [InlineKeyboardButton("سبتمبر", callback_data="month:9"), InlineKeyboardButton("أكتوبر", callback_data="month:10")],
        [InlineKeyboardButton("نوفمبر", callback_data="month:11"), InlineKeyboardButton("ديسمبر", callback_data="month:12")],
        [InlineKeyboardButton("📄 الكل", callback_data="month:all")],
        [InlineKeyboardButton("🔙 رجوع", callback_data="back:year")],
        [InlineKeyboardButton("❌ إلغاء", callback_data="abort")],
    ]
    return InlineKeyboardMarkup(months)


def _confirm_kb(show_back=True):
    buttons = [
        [InlineKeyboardButton("🌐 HTML (موصى به - يعمل على الجوال واللابتوب)", callback_data="export:html")],
        [
            InlineKeyboardButton("📕 PDF", callback_data="export:pdf"),
            InlineKeyboardButton("📗 Excel", callback_data="export:excel"),
        ],
        [InlineKeyboardButton("📘 Word", callback_data="export:word")],
    ]
    if show_back:
        buttons.append([InlineKeyboardButton("🔙 رجوع", callback_data="back:confirm")])
    buttons.append([InlineKeyboardButton("❌ إلغاء", callback_data="abort")])
    return InlineKeyboardMarkup(buttons)


def _action_types_kb():
    """لوحة اختيار نوع الإجراء الطبي"""
    actions = [
        ("📋 كل الإجراءات", "action_type:all"),
        ("🆕 استشارة جديدة", "action_type:استشارة جديدة"),
        ("⚕️ استشارة مع قرار عملية", "action_type:استشارة مع قرار عملية"),
        ("🏁 استشارة أخيرة", "action_type:استشارة أخيرة"),
        ("🚨 طوارئ", "action_type:طوارئ"),
        ("🏥 متابعة في الرقود", "action_type:متابعة في الرقود"),
        ("🔄 مراجعة / عودة دورية", "action_type:مراجعة / عودة دورية"),
        ("🔪 عملية", "action_type:عملية"),
        ("💪 علاج طبيعي", "action_type:علاج طبيعي وإعادة تأهيل"),
        ("🛏️ ترقيد", "action_type:ترقيد"),
        ("🚪 خروج من المستشفى", "action_type:خروج من المستشفى"),
        ("🔬 أشعة وفحوصات", "action_type:أشعة وفحوصات"),
        ("📅 تأجيل موعد", "action_type:تأجيل موعد"),
        ("☢️ جلسة إشعاعي", "action_type:جلسة إشعاعي"),
    ]

    keyboard = []
    # "كل الإجراءات" في صف منفصل
    keyboard.append([InlineKeyboardButton(actions[0][0], callback_data=actions[0][1])])
    # باقي الإجراءات: زرين في كل صف
    for i in range(1, len(actions), 2):
        row = [InlineKeyboardButton(actions[i][0], callback_data=actions[i][1])]
        if i + 1 < len(actions):
            row.append(InlineKeyboardButton(actions[i + 1][0], callback_data=actions[i + 1][1]))
        keyboard.append(row)

    keyboard.append([InlineKeyboardButton("🔙 رجوع", callback_data="back:patient_list")])
    keyboard.append([InlineKeyboardButton("❌ إلغاء", callback_data="abort")])
    return InlineKeyboardMarkup(keyboard)


@admin_handler
async def start_reports_filter(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    if not is_admin(user.id):
        return ConversationHandler.END
    context.user_data.clear()
    await update.message.reply_text("🖨️ اختر نوع الفلترة:", reply_markup=_filters_kb())
    return SELECT_FILTER

@admin_handler
async def handle_filter_choice(update: Update, context: ContextTypes.DEFAULT_TYPE):
    q = update.callback_query
    await q.answer()
    choice = q.data.split(":")[1]
    context.user_data["filter_type"] = choice

    if choice == "patient":
        # ✅ عرض قائمة المرضى للاختيار منها
        context.user_data["mode"] = "print_patient"
        context.user_data["print_patient_page"] = 0  # الصفحة الأولى
        # ✅ مسح أي بيانات قديمة متعلقة بالفلترة
        context.user_data.pop("filter_value", None)
        context.user_data.pop("action_type", None)
        
        # جلب قائمة المرضى من قاعدة البيانات
        with SessionLocal() as session:
            patients = session.query(Patient).filter(
                Patient.full_name.isnot(None),
                Patient.full_name != ""
            ).order_by(Patient.full_name).all()
            patients_list = [(p.id, p.full_name) for p in patients]
        
        if not patients_list:
            await q.edit_message_text(
                "⚠️ لا يوجد مرضى في قاعدة البيانات!",
                reply_markup=InlineKeyboardMarkup([
                    [InlineKeyboardButton("🔙 رجوع", callback_data="back:filter")],
                    [InlineKeyboardButton("❌ إلغاء", callback_data="abort")]
                ])
            )
            return SELECT_FILTER
        
        # حفظ قائمة المرضى في السياق
        context.user_data["patients_list"] = patients_list
        
        # عرض الصفحة الأولى
        text, keyboard = _build_patient_list_keyboard(patients_list, page=0)
        await q.edit_message_text(text, reply_markup=keyboard, parse_mode="Markdown")
        return ENTER_NAME
    
    elif choice == "patient_text":
        # ✅ طباعة حسب المريض (نص داخل Telegram)
        context.user_data["print_patient_page"] = 0  # إعادة تعيين الصفحة
        return await show_patient_selection_for_print_local(q, context)
    
    elif choice == "hospital":
        # ✅ جلب قائمة المستشفيات من ملف doctors.txt
        hospitals_set = set()
        
        try:
            with open("data/doctors.txt", "r", encoding="utf-8") as f:
                for line in f:
                    line = line.strip()
                    if not line or line.startswith("#"):
                        continue
                    parts = [p.strip() for p in line.split("|")]
                    if len(parts) >= 2:
                        hospitals_set.add(parts[1])
        except Exception as e:
            print(f"❌ خطأ في قراءة المستشفيات: {e}")
        
        hospitals_list = sorted(list(hospitals_set))
        
        if not hospitals_list:
            await q.edit_message_text(
                "⚠️ لا توجد مستشفيات!\n\n"
                "تحقق من ملف doctors.txt",
                reply_markup=InlineKeyboardMarkup([
                    [InlineKeyboardButton("🔙 رجوع", callback_data="back:filter")],
                    [InlineKeyboardButton("❌ إلغاء", callback_data="abort")]
                ])
            )
            return SELECT_FILTER
        
        # بناء أزرار المستشفيات
        keyboard = []
        for hospital in hospitals_list:
            # استخدام اسم مختصر للـ callback_data لتجنب تجاوز الحد
            short_name = hospital[:50] if len(hospital) > 50 else hospital
            keyboard.append([InlineKeyboardButton(
                f"🏥 {hospital}",
                callback_data=f"select_hospital:{short_name}"
            )])
        
        # أزرار التحكم
        keyboard.append([InlineKeyboardButton("✍️ كتابة يدوياً", callback_data="hospital:manual")])
        keyboard.append([InlineKeyboardButton("🔙 رجوع", callback_data="back:filter")])
        keyboard.append([InlineKeyboardButton("❌ إلغاء", callback_data="abort")])
        
        await q.edit_message_text(
            "🏥 **اختر المستشفى من القائمة المحدثة:**\n\n"
            f"📊 وجدنا {len(hospitals_list)} مستشفى\n\n"
            "أو اكتب اسم المستشفى يدوياً:",
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode="Markdown"
        )
        return ENTER_NAME
    
    elif choice == "department":
        # جلب قائمة الأقسام من قاعدة البيانات
        with SessionLocal() as session:
            departments = session.query(Department).order_by(Department.name).all()
            dept_names = sorted(list(set([d.name for d in departments if d.name])))
            
            if not dept_names:
                await q.edit_message_text(
                    "⚠️ لا توجد أقسام!\n\n"
                    "لم يتم العثور على أقسام في قاعدة البيانات.",
                    reply_markup=InlineKeyboardMarkup([
                        [InlineKeyboardButton("🔙 رجوع", callback_data="back:filter")],
                        [InlineKeyboardButton("❌ إلغاء", callback_data="abort")]
                    ])
                )
                return SELECT_FILTER
            
            # بناء أزرار الأقسام
            keyboard = []
            for dept in dept_names[:20]:  # أول 20 قسم
                keyboard.append([InlineKeyboardButton(
                    f"🏢 {dept}",
                    callback_data=f"select_dept:{dept[:40]}"  # تقصير للـ callback
                )])
            
            keyboard.append([InlineKeyboardButton("✍️ كتابة يدوياً", callback_data="dept:manual")])
            keyboard.append([InlineKeyboardButton("🔙 رجوع", callback_data="back:filter")])
            keyboard.append([InlineKeyboardButton("❌ إلغاء", callback_data="abort")])
            
            await q.edit_message_text(
                f"🏢 **اختر القسم من القائمة:**\n\n"
                f"📊 وجدنا {len(dept_names)} قسماً\n\n"
                f"أو اكتب اسم القسم يدوياً:",
                reply_markup=InlineKeyboardMarkup(keyboard),
                parse_mode="Markdown"
            )
            return ENTER_NAME
    
    elif choice == "date":
        await q.edit_message_text("📅 اختر السنة:", reply_markup=_years_kb())
        return SELECT_YEAR
    
    else:  # all
        context.user_data["filter_value"] = None
        context.user_data["year_value"] = None
        context.user_data["month_value"] = None
        await q.edit_message_text(
            "✅ **طباعة كل التقارير**\n\n"
            "📋 اختر صيغة التصدير:",
            reply_markup=_confirm_kb(),
            parse_mode="Markdown"
        )
        return CONFIRM_EXPORT

async def get_name(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text.strip()
    if text.lower() in ['إلغاء', 'الغاء', 'cancel']:
        return await cancel_text(update, context)
    
    context.user_data["filter_value"] = text
    
    filter_type = context.user_data.get("filter_type")
    
    # إذا كان المستشفى، نسأل عن القسم أولاً
    if filter_type == "hospital":
        
        keyboard = InlineKeyboardMarkup([
            [InlineKeyboardButton("🏥 المستشفى كاملاً", callback_data="dept_option:all")],
            [InlineKeyboardButton("🏷️ قسم معين فقط", callback_data="dept_option:specific")],
            [InlineKeyboardButton("🔙 رجوع", callback_data="back:filter")],
            [InlineKeyboardButton("❌ إلغاء", callback_data="abort")],
        ])
        
        await update.message.reply_text(
            f"🏥 **تم اختيار المستشفى:** {text}\n\n"
            f"📋 هل تريد طباعة:\n\n"
            f"🏥 **المستشفى كاملاً** (كل الأقسام)\n"
            f"🏷️ **قسم معين فقط** من المستشفى",
            reply_markup=keyboard,
            parse_mode="Markdown"
        )
        
        return SELECT_DEPARTMENT_OPTION
    
    # للمريض أو القسم أو أي نوع آخر
    if filter_type == "patient":
        name_type = "المريض"
        icon = "👤"
    elif filter_type == "department":
        name_type = "القسم"
        icon = "🏢"
    else:
        name_type = "الاسم"
        icon = "📝"
    
    # إنشاء لوحة مفاتيح للخيارات
    keyboard = InlineKeyboardMarkup([
        [InlineKeyboardButton("📅 إضافة فلترة حسب التاريخ", callback_data="add_date_filter:yes")],
        [InlineKeyboardButton("✅ طباعة مباشرة (كل التواريخ)", callback_data="add_date_filter:no")],
        [InlineKeyboardButton("🔙 رجوع", callback_data="back:filter")],
        [InlineKeyboardButton("❌ إلغاء", callback_data="abort")],
    ])
    
    await update.message.reply_text(
        f"{icon} **تم اختيار {name_type}:** {text}\n\n"
        f"📋 هل تريد إضافة فلترة حسب التاريخ أيضاً؟\n\n"
        f"💡 **اختر:**\n"
        f"• إضافة فلترة بالتاريخ (سنة/شهر محدد)\n"
        f"• أو طباعة مباشرة لكل التواريخ",
        reply_markup=keyboard,
        parse_mode="Markdown"
    )
    
    return SELECT_YEAR

async def handle_department_option(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """معالجة اختيار القسم (كامل أو معين)"""
    q = update.callback_query
    await q.answer()
    
    if q.data == "abort":
        await q.edit_message_text("❌ تم إلغاء المحادثة.")
        return ConversationHandler.END
    
    if q.data == "back:filter":
        await q.edit_message_text("🖨️ اختر نوع الفلترة:", reply_markup=_filters_kb())
        return SELECT_FILTER
    
    option = q.data.split(":")[1]
    
    if option == "all":
        # المستشفى كاملاً - نمر لخيارات التاريخ
        context.user_data["department_value"] = None
        
        hospital_name = context.user_data.get("filter_value")
        
        keyboard = InlineKeyboardMarkup([
            [InlineKeyboardButton("📅 إضافة فلترة حسب التاريخ", callback_data="add_date_filter:yes")],
            [InlineKeyboardButton("✅ طباعة مباشرة (كل التواريخ)", callback_data="add_date_filter:no")],
            [InlineKeyboardButton("🔙 رجوع", callback_data="back:dept_option")],
            [InlineKeyboardButton("❌ إلغاء", callback_data="abort")],
        ])
        
        await q.edit_message_text(
            f"🏥 **المستشفى:** {hospital_name} (كامل)\n\n"
            f"📋 هل تريد إضافة فلترة حسب التاريخ أيضاً؟",
            reply_markup=keyboard,
            parse_mode="Markdown"
        )
        
        return SELECT_YEAR
    
    else:  # specific
        # قسم معين - عرض قائمة الأقسام
        hospital_name = context.user_data.get("filter_value")
        
        # جلب قائمة الأقسام من قاعدة البيانات
        with SessionLocal() as s:
            # جلب الأقسام المرتبطة بهذا المستشفى
            hospital = s.query(Hospital).filter(Hospital.name.ilike(f"%{hospital_name}%")).first()
            
            if hospital:
                # جلب التقارير من هذا المستشفى
                dept_ids = s.query(Report.department_id).filter(
                    Report.hospital_id == hospital.id,
                    Report.department_id.isnot(None)
                ).distinct().all()
                
                departments = []
                for (dept_id,) in dept_ids:
                    dept = s.get(Department, dept_id)
                    if dept:
                        departments.append(dept)
            else:
                # إذا لم نجد المستشفى، نجلب كل الأقسام
                departments = s.query(Department).order_by(Department.name).all()
            
            if not departments:
                keyboard = InlineKeyboardMarkup([
                    [InlineKeyboardButton("🔙 رجوع", callback_data="back:dept_option")],
                    [InlineKeyboardButton("❌ إلغاء", callback_data="abort")]
                ])
                
                await q.edit_message_text(
                    f"⚠️ لا توجد أقسام مسجلة في المستشفى!\n\n"
                    f"🏥 المستشفى: {hospital_name}",
                    reply_markup=keyboard,
                    parse_mode="Markdown"
                )
                return SELECT_DEPARTMENT_OPTION
            
            # بناء أزرار الأقسام
            keyboard = []
            for dept in departments[:20]:  # أول 20 قسم
                keyboard.append([InlineKeyboardButton(
                    f"🏷️ {dept.name}",
                    callback_data=f"select_dept:{dept.name}"
                )])
            
            # أزرار التحكم
            keyboard.append([InlineKeyboardButton("✍️ كتابة يدوياً", callback_data="dept:manual")])
            keyboard.append([InlineKeyboardButton("🔙 رجوع", callback_data="back:dept_option")])
            keyboard.append([InlineKeyboardButton("❌ إلغاء", callback_data="abort")])
            
            await q.edit_message_text(
                f"🏥 **المستشفى:** {hospital_name}\n\n"
                f"🏷️ **اختر القسم من القائمة:**\n\n"
                f"📊 وجدنا {len(departments)} قسم\n\n"
                f"أو اكتب يدوياً للبحث",
                reply_markup=InlineKeyboardMarkup(keyboard),
                parse_mode="Markdown"
            )
        
        return ENTER_DEPARTMENT

async def handle_department_name(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """معالجة إدخال اسم القسم"""
    text = update.message.text.strip()
    if text.lower() in ['إلغاء', 'الغاء', 'cancel']:
        return await cancel_text(update, context)
    
    context.user_data["department_value"] = text
    
    hospital_name = context.user_data.get("filter_value")
    
    keyboard = InlineKeyboardMarkup([
        [InlineKeyboardButton("📅 إضافة فلترة حسب التاريخ", callback_data="add_date_filter:yes")],
        [InlineKeyboardButton("✅ طباعة مباشرة (كل التواريخ)", callback_data="add_date_filter:no")],
        [InlineKeyboardButton("🔙 رجوع", callback_data="back:dept_option")],
        [InlineKeyboardButton("❌ إلغاء", callback_data="abort")],
    ])
    
    await update.message.reply_text(
        f"🏥 **المستشفى:** {hospital_name}\n"
        f"🏷️ **القسم:** {text}\n\n"
        f"📋 هل تريد إضافة فلترة حسب التاريخ أيضاً؟",
        reply_markup=keyboard,
        parse_mode="Markdown"
    )
    
    return SELECT_YEAR

async def handle_year_choice(update: Update, context: ContextTypes.DEFAULT_TYPE):
    q = update.callback_query
    await q.answer()
    
    if q.data == "abort":
        await q.edit_message_text("❌ تم إلغاء المحادثة.")
        return ConversationHandler.END
    
    # معالجة زر الرجوع
    if q.data == "back:filter":
        # الرجوع لاختيار نوع الفلترة
        await q.edit_message_text("🖨️ اختر نوع الفلترة:", reply_markup=_filters_kb())
        return SELECT_FILTER

    # الرجوع لاختيار نوع الإجراء (من خطوة التاريخ)
    if q.data == "back:action_type":
        patient_name = context.user_data.get("filter_value", "—")
        await q.edit_message_text(
            f"👤 **المريض:** {patient_name}\n\n"
            f"📌 **اختر نوع الإجراء الطبي:**",
            reply_markup=_action_types_kb(),
            parse_mode="Markdown"
        )
        return SELECT_ACTION_TYPE

    # معالجة زر الرجوع لخيار القسم
    if q.data == "back:dept_option":
        hospital_name = context.user_data.get("filter_value")
        
        keyboard = InlineKeyboardMarkup([
            [InlineKeyboardButton("🏥 المستشفى كاملاً", callback_data="dept_option:all")],
            [InlineKeyboardButton("🏷️ قسم معين فقط", callback_data="dept_option:specific")],
            [InlineKeyboardButton("🔙 رجوع", callback_data="back:filter")],
            [InlineKeyboardButton("❌ إلغاء", callback_data="abort")],
        ])
        
        await q.edit_message_text(
            f"🏥 **تم اختيار المستشفى:** {hospital_name}\n\n"
            f"📋 هل تريد طباعة:\n\n"
            f"🏥 **المستشفى كاملاً** (كل الأقسام)\n"
            f"🏷️ **قسم معين فقط** من المستشفى",
            reply_markup=keyboard,
            parse_mode="Markdown"
        )
        
        return SELECT_DEPARTMENT_OPTION
    
    # معالجة زر الرجوع من التاريخ لقائمة الأقسام
    if q.data == "back:dept_list":
        # إعادة عرض قائمة الأقسام
        hospital_name = context.user_data.get("filter_value")
        
        with SessionLocal() as s:
            hospital = s.query(Hospital).filter(Hospital.name.ilike(f"%{hospital_name}%")).first()
            
            if hospital:
                dept_ids = s.query(Report.department_id).filter(
                    Report.hospital_id == hospital.id,
                    Report.department_id.isnot(None)
                ).distinct().all()
                
                departments = []
                for (dept_id,) in dept_ids:
                    dept = s.get(Department, dept_id)
                    if dept:
                        departments.append(dept)
            else:
                departments = s.query(Department).order_by(Department.name).all()
            
            keyboard = []
            for dept in departments[:20]:
                keyboard.append([InlineKeyboardButton(
                    f"🏷️ {dept.name}",
                    callback_data=f"select_dept:{dept.name}"
                )])
            
            keyboard.append([InlineKeyboardButton("✍️ كتابة يدوياً", callback_data="dept:manual")])
            keyboard.append([InlineKeyboardButton("🔙 رجوع", callback_data="back:dept_option")])
            keyboard.append([InlineKeyboardButton("❌ إلغاء", callback_data="abort")])
            
            await q.edit_message_text(
                f"🏥 **المستشفى:** {hospital_name}\n\n"
                f"🏷️ **اختر القسم من القائمة:**\n\n"
                f"📊 وجدنا {len(departments)} قسم\n\n"
                f"أو اكتب يدوياً للبحث",
                reply_markup=InlineKeyboardMarkup(keyboard),
                parse_mode="Markdown"
            )
        
        return ENTER_DEPARTMENT
    
    # معالجة خيار إضافة فلترة التاريخ
    if q.data.startswith("add_date_filter:"):
        add_date = q.data.split(":")[1]
        
        if add_date == "no":
            # طباعة مباشرة بدون فلترة تاريخ
            context.user_data["year_value"] = None
            context.user_data["month_value"] = None

            filter_type = context.user_data.get("filter_type")
            filter_value = context.user_data.get("filter_value")
            act_type = context.user_data.get("action_type")
            act_label = "كل الإجراءات" if (not act_type or act_type == "all") else act_type

            if filter_type == "patient":
                confirm_msg = f"👤 **المريض:** {filter_value}\n📌 **الإجراء:** {act_label}\n📅 **التاريخ:** كل التواريخ\n\n📋 اختر صيغة التصدير:"
            elif filter_type == "hospital":
                dept_value = context.user_data.get("department_value")
                if dept_value:
                    confirm_msg = f"🏥 **المستشفى:** {filter_value}\n🏷️ **القسم:** {dept_value}\n📅 **التاريخ:** كل التواريخ\n\n📋 اختر صيغة التصدير:"
                else:
                    confirm_msg = f"🏥 **المستشفى:** {filter_value} (كامل)\n📅 **التاريخ:** كل التواريخ\n\n📋 اختر صيغة التصدير:"
            else:
                confirm_msg = "📋 اختر صيغة التصدير:"
            
            await q.edit_message_text(confirm_msg, reply_markup=_confirm_kb(), parse_mode="Markdown")
            return CONFIRM_EXPORT
        else:
            # إضافة فلترة تاريخ
            await q.edit_message_text("📅 اختر السنة:", reply_markup=_years_kb())
            return SELECT_YEAR
    
    year_choice = q.data.split(":")[1]
    context.user_data["year_value"] = year_choice
    
    if year_choice == "all":
        # إذا اختار "الكل" للسنة، انتقل مباشرة للتأكيد
        context.user_data["month_value"] = None

        filter_type = context.user_data.get("filter_type")
        filter_value = context.user_data.get("filter_value")
        act_type = context.user_data.get("action_type")
        act_label = "كل الإجراءات" if (not act_type or act_type == "all") else act_type

        if filter_type == "patient":
            confirm_msg = f"👤 **المريض:** {filter_value}\n📌 **الإجراء:** {act_label}\n📅 **التاريخ:** كل السنوات\n\n📋 اختر صيغة التصدير:"
        elif filter_type == "hospital":
            dept_value = context.user_data.get("department_value")
            if dept_value:
                confirm_msg = f"🏥 **المستشفى:** {filter_value}\n🏷️ **القسم:** {dept_value}\n📅 **التاريخ:** كل السنوات\n\n📋 اختر صيغة التصدير:"
            else:
                confirm_msg = f"🏥 **المستشفى:** {filter_value} (كامل)\n📅 **التاريخ:** كل السنوات\n\n📋 اختر صيغة التصدير:"
        else:
            confirm_msg = "✅ تم اختيار جميع السنوات. هل تريد الطباعة؟"
        
        await q.edit_message_text(confirm_msg, reply_markup=_confirm_kb(), parse_mode="Markdown")
        return CONFIRM_EXPORT
    else:
        # إذا اختار سنة محددة، اعرض الشهور
        await q.edit_message_text(f"📅 اختر الشهر للسنة {year_choice}:", reply_markup=_months_kb())
        return SELECT_MONTH

async def handle_month_choice(update: Update, context: ContextTypes.DEFAULT_TYPE):
    q = update.callback_query
    await q.answer()
    
    if q.data == "abort":
        await q.edit_message_text("❌ تم إلغاء المحادثة.")
        return ConversationHandler.END
    
    # معالجة زر الرجوع
    if q.data == "back:year":
        # الرجوع لاختيار السنة
        await q.edit_message_text("📅 اختر السنة:", reply_markup=_years_kb())
        return SELECT_YEAR
    
    month_choice = q.data.split(":")[1]
    context.user_data["month_value"] = month_choice
    
    year = context.user_data.get("year_value", "غير محدد")
    filter_type = context.user_data.get("filter_type")
    filter_value = context.user_data.get("filter_value")
    
    month_names = {
        "1": "يناير", "2": "فبراير", "3": "مارس", "4": "أبريل",
        "5": "مايو", "6": "يونيو", "7": "يوليو", "8": "أغسطس",
        "9": "سبتمبر", "10": "أكتوبر", "11": "نوفمبر", "12": "ديسمبر"
    }
    
    if month_choice == "all":
        month_text = "جميع الشهور"
    else:
        month_text = month_names.get(month_choice, month_choice)
    
    # بناء رسالة التأكيد
    act_type = context.user_data.get("action_type")
    act_label = "كل الإجراءات" if (not act_type or act_type == "all") else act_type

    if filter_type == "patient" and filter_value:
        confirm_msg = f"👤 **المريض:** {filter_value}\n📌 **الإجراء:** {act_label}\n📅 **التاريخ:** {year} - {month_text}\n\n📋 اختر صيغة التصدير:"
    elif filter_type == "hospital" and filter_value:
        dept_value = context.user_data.get("department_value")
        if dept_value:
            confirm_msg = f"🏥 **المستشفى:** {filter_value}\n🏷️ **القسم:** {dept_value}\n📅 **التاريخ:** {year} - {month_text}\n\n📋 اختر صيغة التصدير:"
        else:
            confirm_msg = f"🏥 **المستشفى:** {filter_value} (كامل)\n📅 **التاريخ:** {year} - {month_text}\n\n📋 اختر صيغة التصدير:"
    else:
        confirm_msg = f"📅 **التاريخ:** {year} - {month_text}\n\n📋 اختر صيغة التصدير:"
    
    await q.edit_message_text(confirm_msg, reply_markup=_confirm_kb(), parse_mode="Markdown")
    return CONFIRM_EXPORT

def _query_reports(filter_type, name_val, year_val, month_val, dept_val=None, action_type=None):
    with SessionLocal() as s:
        base_query = s.query(Report)

        if filter_type == "patient" and name_val:
            # استخدام ON clause صريح
            base_query = base_query.join(Patient, Report.patient_id == Patient.id).filter(Patient.full_name.ilike(f"%{name_val}%"))
        elif filter_type == "hospital" and name_val:
            # استخدام ON clause صريح
            # تطبيع اسم المستشفى للبحث المرن (استبدال الفواصل والشرطات)
            search_name = name_val.replace(",", "%").replace(" - ", "%").replace("-", "%")
            base_query = base_query.join(Hospital, Report.hospital_id == Hospital.id).filter(Hospital.name.ilike(f"%{search_name}%"))
            
            # إذا كان هناك فلتر قسم أيضاً
            if dept_val:
                base_query = base_query.join(Department, Report.department_id == Department.id).filter(Department.name.ilike(f"%{dept_val}%"))
        elif filter_type == "department" and name_val:
            # فلترة حسب القسم فقط - استخدام ON clause صريح
            base_query = base_query.join(Department, Report.department_id == Department.id).filter(Department.name.ilike(f"%{name_val}%"))

        # فلترة حسب نوع الإجراء الطبي
        if action_type and action_type != "all":
            base_query = base_query.filter(Report.medical_action == action_type)

        # فلترة حسب السنة والشهر
        if year_val and year_val != "all":
            try:
                year = int(year_val)
                if month_val and month_val != "all":
                    # سنة وشهر محددين
                    month = int(month_val)
                    start = datetime(year, month, 1)
                    end = datetime(year + 1, 1, 1) if month == 12 else datetime(year, month + 1, 1)
                    base_query = base_query.filter(Report.report_date >= start, Report.report_date < end)
                else:
                    # سنة كاملة
                    start = datetime(year, 1, 1)
                    end = datetime(year + 1, 1, 1)
                    base_query = base_query.filter(Report.report_date >= start, Report.report_date < end)
            except Exception:
                return "FORMAT_ERR", [], None, None
        else:
            # ✅ إذا اختار "الكل" - من 2025 فصاعداً
            start = datetime(2025, 1, 1)
            base_query = base_query.filter(Report.report_date >= start)

        reports = base_query.order_by(Report.report_date.desc()).all()
        
        # إنشاء الإحصائيات والرسوم البيانية
        stats = generate_statistics(reports, s) if reports else {}
        charts = create_charts(stats, filter_type) if stats else {}

        rows = []
        for r in reports:
            # حاول استخدام الحقول المكررة داخل التقرير أولاً (denormalized fields)
            p_name = (r.patient_name.strip() if getattr(r, 'patient_name', None) else None) or None
            h_name = (r.hospital_name.strip() if getattr(r, 'hospital_name', None) else None) or None
            dept_name = (r.department.strip() if getattr(r, 'department', None) else None) or None
            doc_name = (r.doctor_name.strip() if getattr(r, 'doctor_name', None) else None) or None

            # إذا كانت الحقول المكررة غير متوفرة، حاول جلبها من الجداول المرتبطة عبر المفاتيح الخارجية
            try:
                if not p_name and r.patient_id:
                    p = s.get(Patient, r.patient_id)
                    p_name = p.full_name if p and p.full_name else None
                if not h_name and r.hospital_id:
                    h = s.get(Hospital, r.hospital_id)
                    h_name = h.name if h and h.name else None
                if not dept_name and r.department_id:
                    d = s.get(Department, r.department_id)
                    dept_name = d.name if d and d.name else None
                if not doc_name and r.doctor_id:
                    doc = s.get(Doctor, r.doctor_id)
                    # Doctor model may have 'full_name' or 'name'
                    doc_name = (doc.full_name if getattr(doc, 'full_name', None) else getattr(doc, 'name', None)) if doc else None
            except Exception:
                pass

            rows.append({
                "id": r.id,
                "report_date": r.report_date.strftime("%Y-%m-%d %H:%M") if r.report_date else "",
                "patient_name": p_name or "—",
                "hospital": h_name or "—",
                "department": dept_name or "—",
                "doctor": doc_name or "—",
                "complaint": r.complaint_text or "—",
                "action": r.medical_action or "—",
                "decision": r.doctor_decision or "—",
                "followup_date": r.followup_date.strftime("%Y-%m-%d") if r.followup_date else "—",
                "followup_reason": r.followup_reason or "—",
            })
        return None, rows, stats, charts


def _create_reports_charts_sync(hospitals_counter, departments_counter, actions_counter, rows):
    """إنشاء رسوم بيانية للتقارير"""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    from matplotlib import font_manager
    import base64
    from io import BytesIO
    from arabic_reshaper import reshape
    from bidi.algorithm import get_display
    from collections import defaultdict
    
    # إعداد الخط العربي
    try:
        font_path = "C:\\Windows\\Fonts\\tahoma.ttf"
        import os
        if os.path.exists(font_path):
            prop = font_manager.FontProperties(fname=font_path)
            plt.rcParams['font.family'] = prop.get_name()
    except:
        pass
    
    plt.rcParams['axes.unicode_minus'] = False
    
    def ar(text):
        try:
            return get_display(reshape(str(text)))
        except:
            return str(text)
    
    def fig_to_base64(fig):
        buffer = BytesIO()
        fig.savefig(buffer, format='png', dpi=150, bbox_inches='tight', facecolor='white')
        buffer.seek(0)
        image_base64 = base64.b64encode(buffer.read()).decode()
        plt.close(fig)
        return image_base64
    
    charts = {}
    
    try:
        # 1. رسم دائري للمستشفيات
        if hospitals_counter and len(hospitals_counter) > 0:
            fig, ax = plt.subplots(figsize=(10, 7))
            hospitals = list(hospitals_counter.keys())[:8]
            values = list(hospitals_counter.values())[:8]
            colors = ['#3498db', '#e74c3c', '#2ecc71', '#f39c12', '#9b59b6', '#1abc9c', '#e67e22', '#34495e']
            
            wedges, texts, autotexts = ax.pie(
                values,
                labels=[ar(h) for h in hospitals],
                autopct='%1.1f%%',
                startangle=90,
                colors=colors,
                textprops={'fontsize': 12, 'weight': 'bold'}
            )
            
            for autotext in autotexts:
                autotext.set_color('white')
                autotext.set_fontsize(11)
            
            ax.set_title(ar('توزيع التقارير حسب المستشفيات'), fontsize=16, weight='bold', pad=20)
            charts['hospitals_chart'] = fig_to_base64(fig)
        
        # 2. رسم أعمدة للأقسام
        if departments_counter and len(departments_counter) > 0:
            fig, ax = plt.subplots(figsize=(12, 7))
            depts = list(departments_counter.keys())[:10]
            values = list(departments_counter.values())[:10]
            colors_gradient = ['#e74c3c', '#e67e22', '#f39c12', '#f1c40f', '#2ecc71', '#1abc9c', '#3498db', '#9b59b6', '#e84393', '#fd79a8']
            
            bars = ax.barh(range(len(depts)), values, color=colors_gradient[:len(depts)], edgecolor='white', linewidth=2)
            
            ax.set_yticks(range(len(depts)))
            ax.set_yticklabels([ar(d) for d in depts], fontsize=11)
            ax.set_xlabel(ar('عدد التقارير'), fontsize=13, weight='bold')
            ax.set_title(ar('توزيع التقارير حسب الأقسام'), fontsize=16, weight='bold', pad=20)
            ax.grid(axis='x', alpha=0.3, linestyle='--')
            
            for i, v in enumerate(values):
                ax.text(v + 0.5, i, str(v), va='center', fontsize=10, weight='bold')
            
            charts['departments_chart'] = fig_to_base64(fig)
        
        # 3. رسم أعمدة للإجراءات
        if actions_counter and len(actions_counter) > 0:
            fig, ax = plt.subplots(figsize=(12, 7))
            actions = list(actions_counter.keys())[:8]
            values = list(actions_counter.values())[:8]
            colors_gradient = ['#e74c3c', '#f39c12', '#2ecc71', '#3498db', '#9b59b6', '#1abc9c', '#e67e22', '#34495e']
            
            bars = ax.bar(range(len(actions)), values, color=colors_gradient[:len(actions)], edgecolor='white', linewidth=2)
            
            ax.set_xticks(range(len(actions)))
            ax.set_xticklabels([ar(a) for a in actions], fontsize=10, rotation=30, ha='right')
            ax.set_ylabel(ar('عدد التقارير'), fontsize=13, weight='bold')
            ax.set_title(ar('توزيع الإجراءات الطبية'), fontsize=16, weight='bold', pad=20)
            ax.grid(axis='y', alpha=0.3, linestyle='--')
            
            for i, v in enumerate(values):
                ax.text(i, v + 0.5, str(v), ha='center', fontsize=10, weight='bold')
            
            charts['actions_chart'] = fig_to_base64(fig)
        
        # 4. رسم خطي للتطور الزمني
        try:
            date_counter = defaultdict(int)
            for r in rows:
                date_str = r.get('report_date', '')
                if date_str:
                    try:
                        date_only = date_str.split()[0] if ' ' in date_str else date_str
                        date_counter[date_only] += 1
                    except:
                        pass
            
            if date_counter and len(date_counter) > 1:
                fig, ax = plt.subplots(figsize=(14, 7))
                sorted_dates = sorted(date_counter.items())
                dates = [d[0] for d in sorted_dates]
                counts = [d[1] for d in sorted_dates]
                
                ax.plot(range(len(dates)), counts, marker='o', linewidth=2.5, markersize=8, 
                       color='#3498db', markerfacecolor='#e74c3c', markeredgewidth=2, markeredgecolor='white')
                ax.fill_between(range(len(dates)), counts, alpha=0.3, color='#3498db')
                
                ax.set_xticks(range(len(dates)))
                ax.set_xticklabels(dates, fontsize=9, rotation=45, ha='right')
                ax.set_ylabel(ar('عدد التقارير'), fontsize=13, weight='bold')
                ax.set_title(ar('تطور عدد التقارير عبر الزمن'), fontsize=16, weight='bold', pad=20)
                ax.grid(True, alpha=0.3, linestyle='--')
                
                for i, v in enumerate(counts):
                    ax.text(i, v + 0.2, str(v), ha='center', fontsize=9, weight='bold')
                
                charts['timeline_chart'] = fig_to_base64(fig)
        except Exception as e:
            logger.warning(f"⚠️ تعذر إنشاء رسم الزمن: {e}")
        
    except Exception as e:
        logger.error(f"❌ خطأ في إنشاء الرسوم البيانية: {e}")
    
    return charts

def _generate_reports_pdf_with_charts_sync(pdf_data, charts_data):
    """إنشاء PDF بتصميم فريد مع الرسوم البيانية"""
    from datetime import datetime
    import sys
    import os
    
    try:
        # إنشاء HTML مباشرة بدون قالب
        generated_at = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        date_from = pdf_data.get('date_from', 'غير محدد')
        date_to = pdf_data.get('date_to', 'غير محدد')
        total_reports = pdf_data.get('total_reports', 0)
        total_patients = pdf_data.get('total_patients', 0)
        hospitals_count = pdf_data.get('hospitals_count', 0)
        doctors_count = pdf_data.get('doctors_count', 0)
        hospitals_data = pdf_data.get('hospitals_data', [])
        departments_data = pdf_data.get('departments_data', [])
        actions_data = pdf_data.get('actions_data', [])
        charts = charts_data or {}
        
        # بناء HTML مباشرة
        html_content = f'''<!DOCTYPE html>
<html dir="rtl" lang="ar">
<head>
    <meta charset="UTF-8">
    <title>تقرير طبي شامل</title>
    <style>
        body {{ font-family: Arial, Tahoma, sans-serif; direction: rtl; text-align: right; padding: 20px; }}
        h1 {{ color: #2c3e50; }}
        table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: right; }}
        th {{ background-color: #3498db; color: white; }}
    </style>
</head>
<body>
    <h1>📊 التقرير الطبي الشامل</h1>
    <p><strong>تاريخ الإنشاء:</strong> {generated_at}</p>
    <p><strong>الفترة:</strong> {date_from} - {date_to}</p>
    <p><strong>عدد التقارير:</strong> {total_reports}</p>
    <p><strong>عدد المرضى:</strong> {total_patients}</p>
    <p><strong>عدد المستشفيات:</strong> {hospitals_count}</p>
    <p><strong>عدد الأطباء:</strong> {doctors_count}</p>
'''
        
        # إضافة الرسوم البيانية إذا كانت موجودة
        if charts:
            html_content += '<h2>📈 الرسوم البيانية</h2>'
            for chart_name, chart_base64 in charts.items():
                html_content += f'<img src="data:image/png;base64,{chart_base64}" style="max-width: 100%; margin: 10px 0;" />'
        
        html_content += '''
</body>
</html>
'''
        
        # حفظ أو تحويل إلى PDF
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        os.makedirs("exports", exist_ok=True)
        
        if sys.platform.startswith('win'):
            # Windows: حفظ HTML
            html_path = f"exports/reports_print_{timestamp}.html"
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            logger.info(f"✅ تم حفظ HTML: {html_path}")
            return html_path
        else:
            # Linux/Cloud Run: استخدام WeasyPrint
            from weasyprint import HTML
            pdf_path = f"exports/reports_print_{timestamp}.pdf"
            HTML(string=html_content, base_url=os.path.dirname(os.path.abspath(pdf_path))).write_pdf(pdf_path)
            logger.info(f"✅ تم إنشاء PDF: {pdf_path}")
            return pdf_path
            
    except Exception as e:
        logger.error(f"❌ خطأ في إنشاء PDF: {e}", exc_info=True)
        raise

@admin_handler
async def confirm_export(update: Update, context: ContextTypes.DEFAULT_TYPE):
    q = update.callback_query
    if q:
        await q.answer()
        
        # معالجة اختيار الصيغة
        if q.data.startswith("export:") and q.data != "export:yes" and q.data != "export:no":
            export_format = q.data.split(":")[1]
            context.user_data["export_format"] = export_format
            logger.info(f"✅ تم اختيار نوع التصدير: {export_format}")
            # المتابعة مباشرة للتصدير - لا نستدعي confirm_export مرة أخرى
            # سنتابع مباشرة إلى عملية التصدير
        elif q.data == "export:no":
            await q.edit_message_text("🚫 تم الإلغاء.")
            return ConversationHandler.END
        elif q.data == "abort":
            await q.edit_message_text("❌ تم إلغاء المحادثة.")
            return ConversationHandler.END
        elif q.data == "back:confirm":
            f_type = context.user_data.get("filter_type")
            name_val = context.user_data.get("filter_value")
            year_val = context.user_data.get("year_value")
            month_val = context.user_data.get("month_value")
            
            # تحديد إلى أين نرجع
            if f_type in ["patient", "hospital"] and name_val:
                # إذا كان اختار مريض/مستشفى مع تاريخ، يرجع للشهور
                if year_val and year_val != "all":
                    await q.edit_message_text(f"📅 اختر الشهر للسنة {year_val}:", reply_markup=_months_kb())
                    return SELECT_MONTH
                else:
                    # يرجع للسنوات
                    await q.edit_message_text("📅 اختر السنة:", reply_markup=_years_kb())
                    return SELECT_YEAR
            elif f_type == "date":
                # إذا كان اختار تاريخ فقط، يرجع للشهور
                if year_val and month_val:
                    await q.edit_message_text(f"📅 اختر الشهر للسنة {year_val}:", reply_markup=_months_kb())
                    return SELECT_MONTH
                else:
                    await q.edit_message_text("📅 اختر السنة:", reply_markup=_years_kb())
                    return SELECT_YEAR
            else:
                # الرجوع للخيارات الرئيسية
                await q.edit_message_text("🖨️ اختر نوع الفلترة:", reply_markup=_filters_kb())
                return SELECT_FILTER

    f_type = context.user_data.get("filter_type")
    name_val = context.user_data.get("filter_value")
    year_val = context.user_data.get("year_value")
    month_val = context.user_data.get("month_value")

    dept_val = context.user_data.get("department_value")
    action_type = context.user_data.get("action_type")
    fmt_err, rows, stats, charts = _query_reports(f_type, name_val, year_val, month_val, dept_val, action_type)
    if fmt_err == "FORMAT_ERR":
        if q:
            await q.edit_message_text("⚠️ تنسيق الشهر غير صحيح. أرسل مثل: 2025-10 أو 'الكل'")
        else:
            await update.message.reply_text("⚠️ تنسيق الشهر غير صحيح. أرسل مثل: 2025-10 أو 'الكل'")
        return ConversationHandler.END

    if not rows:
        # بناء رسالة واضحة عند عدم وجود تقارير
        if f_type == "patient":
            msg = f"⚠️ **لا توجد تقارير للمريض:** {name_val}"
        elif f_type == "hospital":
            msg = f"⚠️ **لا توجد تقارير للمستشفى:** {name_val}"
        elif f_type == "date":
            if year_val and month_val:
                msg = f"⚠️ **لا توجد تقارير في:** {year_val} - {month_val}"
            else:
                msg = "⚠️ **لا توجد تقارير في الفترة المحددة**"
        else:
            msg = "⚠️ **لا توجد تقارير مطابقة**"
        
        if q:
            await q.edit_message_text(msg, parse_mode="Markdown")
        else:
            await update.message.reply_text(msg, parse_mode="Markdown")
        return ConversationHandler.END

    # بناء عنوان التقرير والبيانات
    action_label = None
    if action_type and action_type != "all":
        action_label = action_type

    if f_type == "patient":
        title = f"تقارير المريض: {name_val}"
        filter_desc = f"المريض: {name_val}"
        if action_label:
            title += f" ({action_label})"
            filter_desc += f" | الإجراء: {action_label}"
    elif f_type == "hospital":
        title = f"تقارير المستشفى: {name_val}"
        filter_desc = f"المستشفى: {name_val}"
    elif f_type == "date":
        title = "تقارير طبية حسب التاريخ"
        filter_desc = "فلترة حسب التاريخ"
    else:
        title = "تقارير طبية شاملة"
        filter_desc = "جميع التقارير"

    # إضافة معلومات التاريخ إذا وجدت
    if year_val and year_val != "all":
        if month_val and month_val != "all":
            month_names = {
                "1": "يناير", "2": "فبراير", "3": "مارس", "4": "أبريل",
                "5": "مايو", "6": "يونيو", "7": "يوليو", "8": "أغسطس",
                "9": "سبتمبر", "10": "أكتوبر", "11": "نوفمبر", "12": "ديسمبر"
            }
            month_name = month_names.get(month_val, month_val)
            filter_desc += f" | {year_val} - {month_name}"
        else:
            filter_desc += f" | سنة {year_val}"
    elif year_val == "all" or not year_val:
        filter_desc += " | كل التواريخ"

    context_data = {
        "title": title,
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "filter_used": f_type or "all",
        "filter_value": filter_desc,
        "reports": rows,
        "stats": stats or {},
        "charts": charts or {},
    }
    
    os.makedirs("exports", exist_ok=True)
    
    try:
        import sys
        
        # الحصول على الصيغة المختارة
        export_format = context.user_data.get("export_format", "pdf")
        format_names = {"pdf": "PDF", "excel": "Excel", "word": "Word", "html": "HTML"}
        format_name = format_names.get(export_format, "PDF")
        
        logger.info(f"🖨️ بدء إنشاء ملف {format_name} - عدد التقارير: {len(rows)}, نوع التصدير: {export_format}")
        
        if q:
            await q.edit_message_text(f"📝 جاري إنشاء ملف {format_name}...")
        else:
            await update.message.reply_text(f"📝 جاري إنشاء ملف {format_name}...")
        
        # اختيار الصيغة
        file_path = None
        
        # using a plain block instead of nested try to keep structure simple
        if True:
            if export_format == "excel":
                logger.info("📗 بدء تصدير Excel...")
                file_path = await asyncio.to_thread(export_to_excel, rows, f"reports_{f_type or 'all'}")
                logger.info(f"✅ تم إنشاء Excel: {file_path}")
            elif export_format == "word":
                logger.info("📘 بدء تصدير Word...")
                file_path = await asyncio.to_thread(export_to_word, rows, f"reports_{f_type or 'all'}")
                logger.info(f"✅ تم إنشاء Word: {file_path}")
            elif export_format == "html":
                logger.info("🌐 بدء تصدير HTML...")
                file_path = await asyncio.to_thread(export_to_html, rows, f"reports_{f_type or 'all'}", filter_type=f_type, title=title, filter_desc=filter_desc)
                logger.info(f"✅ تم إنشاء HTML: {file_path}")
            else:  # pdf (default)
                logger.info("📕 بدء تصدير PDF...")
                # استخدام نفس نظام التحليلات الاحترافي
                from services.pdf_generator_enhanced import generate_data_analysis_pdf_with_tables
                
                # تحضير البيانات بنفس صيغة التحليلات
                from collections import Counter
                
                # تنفيذ عملية تحضير البيانات والرسوم في thread منفصل لتجنب الحظر
                def prepare_pdf_data_and_charts():
                    pdf_data = {
                        'date_from': context_data.get('filter_value', 'غير محدد'),
                        'date_to': datetime.now().strftime('%Y-%m-%d'),
                        'total_reports': len(rows),
                        'total_patients': len(set(r.get('patient_name') for r in rows if r.get('patient_name'))),
                        'hospitals_count': len(set(r.get('hospital_name') for r in rows if r.get('hospital_name'))),
                        'doctors_count': len(set(r.get('doctor_name') for r in rows if r.get('doctor_name'))),
                        'hospitals_data': [],
                        'departments_data': [],
                        'doctors_data': [],
                        'complaints_data': [],
                        'actions_data': [],
                        'top_patients': []
                    }
                    
                    hospitals_counter = Counter()
                    departments_counter = Counter()
                    doctors_counter = Counter()
                    actions_counter = Counter()
                    complaints_counter = Counter()
                    
                    for r in rows:
                        if r.get('hospital_name'):
                            hospitals_counter[r['hospital_name']] += 1
                        if r.get('department_name'):
                            departments_counter[r['department_name']] += 1
                        if r.get('doctor_name'):
                            doctors_counter[r['doctor_name']] += 1
                        if r.get('medical_action'):
                            actions_counter[r['medical_action']] += 1
                        if r.get('complaint_text'):
                            complaints_counter[r['complaint_text'][:50]] += 1
                    
                    pdf_data['hospitals_data'] = [
                        {
                            'name': k, 
                            'reports_count': v,
                            'count': v,
                            'percentage': float(f"{(v/len(rows)*100):.1f}"),
                            'percent': f"{(v/len(rows)*100):.1f}"
                        } 
                        for k, v in hospitals_counter.most_common(10)
                    ]
                    
                    pdf_data['departments_data'] = [
                        {
                            'name': k, 
                            'count': v,
                            'percentage': float(f"{(v/len(rows)*100):.1f}"),
                            'percent': f"{(v/len(rows)*100):.1f}"
                        } 
                        for k, v in departments_counter.most_common(10)
                    ]
                    
                    pdf_data['doctors_data'] = [
                        {
                            'name': k, 
                            'reports_count': v,
                            'count': v,
                            'percentage': float(f"{(v/len(rows)*100):.1f}"),
                            'percent': f"{(v/len(rows)*100):.1f}"
                        } 
                        for k, v in doctors_counter.most_common(10)
                    ]
                    
                    pdf_data['actions_data'] = [
                        {
                            'name': k, 
                            'count': v,
                            'percentage': float(f"{(v/len(rows)*100):.1f}"),
                            'percent': f"{(v/len(rows)*100):.1f}"
                        } 
                        for k, v in actions_counter.most_common(10)
                    ]
                    
                    pdf_data['complaints_data'] = [
                        {
                            'name': k, 
                            'count': v,
                            'percentage': float(f"{(v/len(rows)*100):.1f}"),
                            'percent': f"{(v/len(rows)*100):.1f}"
                        } 
                        for k, v in complaints_counter.most_common(10)
                    ]
                    
                    patients_counter = Counter()
                    patients_last_visit = {}
                    for r in rows:
                        patient_name = r.get('patient_name')
                        if patient_name:
                            patients_counter[patient_name] += 1
                            visit_date = r.get('report_date', 'غير محدد')
                            if patient_name not in patients_last_visit or visit_date > patients_last_visit.get(patient_name, ''):
                                patients_last_visit[patient_name] = visit_date
                    
                    pdf_data['top_patients'] = [
                        {
                            'name': k, 
                            'visits': v,
                            'last_visit': patients_last_visit.get(k, 'غير محدد')
                        } 
                        for k, v in patients_counter.most_common(10)
                    ]
                    
                    # إنشاء الرسوم البيانية (متزامن الآن)
                    charts_data = _create_reports_charts_sync(
                        hospitals_counter, 
                        departments_counter, 
                        actions_counter,
                        rows
                    )
                    
                    # إنشاء PDF (متزامن الآن)
                    return _generate_reports_pdf_with_charts_sync(pdf_data, charts_data)

                file_path = await asyncio.to_thread(prepare_pdf_data_and_charts)
                logger.info(f"✅ تم إنشاء PDF: {file_path}")
            
            if file_path and os.path.exists(file_path):
                if q:
                    await q.get_bot().send_document(
                        chat_id=q.message.chat_id,
                        document=open(file_path, "rb"),
                        caption=f"✅ {format_name} - {len(rows)} تقرير"
                    )
                else:
                    await update.message.reply_document(
                        document=open(file_path, "rb"),
                        caption=f"✅ {format_name} - {len(rows)} تقرير"
                    )
            else:
                raise Exception(f"فشل في إنشاء ملف {format_name}")
            
    except Exception as e:
        import traceback
        tb = traceback.format_exc()
        logger.error(f"❌ حدث خطأ أثناء إنشاء {format_name}: {e}\n{tb}")
        # نرسل مقتطف من التتبع للمستخدم للمساعدة في التشخيص
        short_tb = tb if len(tb) <= 1500 else tb[-1500:]
        error_msg = f"❌ حدث خطأ أثناء إنشاء {format_name}:\n{e}\n\nTraceback (truncated):\n{short_tb}"
        if q:
            try:
                await q.edit_message_text(error_msg)
            except Exception:
                # إذا فشل تعديل الرسالة (مثلاً انتهت صلاحية callback) أرسل رسالة عادية
                await q.get_bot().send_message(chat_id=q.message.chat_id, text=error_msg)
        else:
            await update.message.reply_text(error_msg)

    context.user_data.clear()
    return ConversationHandler.END

async def cancel_text(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data.clear()
    await update.message.reply_text("❌ تم إلغاء المحادثة.")
    return ConversationHandler.END

async def handle_department_selection(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """معالجة اختيار القسم من القائمة"""
    q = update.callback_query
    await q.answer()
    
    dept_name = q.data.split(":", 1)[1]
    context.user_data["filter_value"] = dept_name
    
    # اختيار الصيغة
    await q.edit_message_text(
        f"🏢 **تم اختيار القسم:** {dept_name}\n\n"
        f"📋 اختر صيغة التصدير:",
        reply_markup=_confirm_kb(),
        parse_mode="Markdown"
    )
    
    return CONFIRM_EXPORT

async def handle_hospital_selection(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """معالجة اختيار المستشفى من القائمة"""
    q = update.callback_query
    await q.answer()
    
    hospital_name = q.data.split(":", 1)[1]
    context.user_data["filter_value"] = hospital_name
    
    # الانتقال مباشرة لخيار القسم
    keyboard = InlineKeyboardMarkup([
        [InlineKeyboardButton("🏥 المستشفى كاملاً", callback_data="dept_option:all")],
        [InlineKeyboardButton("🏷️ قسم معين فقط", callback_data="dept_option:specific")],
        [InlineKeyboardButton("🔙 رجوع", callback_data="back:filter")],
        [InlineKeyboardButton("❌ إلغاء", callback_data="abort")],
    ])
    
    await q.edit_message_text(
        f"🏥 **تم اختيار المستشفى:** {hospital_name}\n\n"
        f"📋 هل تريد طباعة:\n\n"
        f"🏥 **المستشفى كاملاً** (كل الأقسام)\n"
        f"🏷️ **قسم معين فقط** من المستشفى",
        reply_markup=keyboard,
        parse_mode="Markdown"
    )
    
    return SELECT_DEPARTMENT_OPTION

async def handle_print_patient_selection(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """معالجة اختيار المريض من القائمة للطباعة"""
    q = update.callback_query
    await q.answer()
    
    patient_id = int(q.data.split(":", 1)[1])
    
    # جلب اسم المريض من قاعدة البيانات
    with SessionLocal() as session:
        patient = session.query(Patient).filter_by(id=patient_id).first()
        if patient:
            patient_name = patient.full_name
        else:
            await q.edit_message_text("❌ لم يتم العثور على المريض!")
            return SELECT_FILTER
    
    context.user_data["filter_value"] = patient_name
    context.user_data["filter_type"] = "patient"

    # ✅ حساب عدد التقارير لكل نوع إجراء لهذا المريض
    with SessionLocal() as session:
        from sqlalchemy import func as sqlfunc
        action_counts = session.query(
            Report.medical_action, sqlfunc.count(Report.id)
        ).filter(
            Report.patient_id == patient_id
        ).group_by(Report.medical_action).all()

        total = sum(c for _, c in action_counts)
        action_summary = "\n".join(
            f"  • {action or 'غير محدد'}: {count}"
            for action, count in sorted(action_counts, key=lambda x: x[1], reverse=True)
        )

    # ✅ الانتقال لاختيار نوع الإجراء الطبي (الخطوة الجديدة)
    await q.edit_message_text(
        f"✅ **تم اختيار المريض:** {patient_name}\n"
        f"📊 **إجمالي التقارير:** {total}\n\n"
        f"📌 **التقارير حسب الإجراء:**\n{action_summary}\n\n"
        f"🔽 **اختر نوع الإجراء:**",
        reply_markup=_action_types_kb(),
        parse_mode="Markdown"
    )
    return SELECT_ACTION_TYPE


async def handle_patient_page(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """معالجة التنقل بين صفحات قائمة المرضى"""
    q = update.callback_query
    await q.answer()
    
    page = int(q.data.split(":", 1)[1])
    # ✅ حفظ رقم الصفحة الحالية في context
    context.user_data["print_patient_page"] = page
    
    patients_list = context.user_data.get("patients_list", [])
    
    if not patients_list:
        # إعادة جلب القائمة من قاعدة البيانات
        with SessionLocal() as session:
            patients = session.query(Patient).filter(
                Patient.full_name.isnot(None),
                Patient.full_name != ""
            ).order_by(Patient.full_name).all()
            patients_list = [(p.id, p.full_name) for p in patients]
            context.user_data["patients_list"] = patients_list
    
    text, keyboard = _build_patient_list_keyboard(patients_list, page=page)
    await q.edit_message_text(text, reply_markup=keyboard, parse_mode="Markdown")
    return ENTER_NAME


async def handle_action_type_choice(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """✅ معالجة اختيار نوع الإجراء الطبي (خطوة جديدة)"""
    q = update.callback_query
    await q.answer()

    data = q.data

    # زر الرجوع لقائمة المرضى
    if data == "back:patient_list":
        patients_list = context.user_data.get("patients_list", [])
        if not patients_list:
            with SessionLocal() as session:
                patients = session.query(Patient).filter(
                    Patient.full_name.isnot(None),
                    Patient.full_name != ""
                ).order_by(Patient.full_name).all()
                patients_list = [(p.id, p.full_name) for p in patients]
                context.user_data["patients_list"] = patients_list

        page = context.user_data.get("print_patient_page", 0)
        text, keyboard = _build_patient_list_keyboard(patients_list, page=page)
        await q.edit_message_text(text, reply_markup=keyboard, parse_mode="Markdown")
        return ENTER_NAME

    # زر الإلغاء
    if data == "abort":
        await q.edit_message_text("❌ تم إلغاء الطباعة.")
        return ConversationHandler.END

    # استخراج نوع الإجراء
    action_type = data.split(":", 1)[1]  # "all" أو اسم الإجراء
    context.user_data["action_type"] = action_type

    patient_name = context.user_data.get("filter_value", "—")
    action_label = "كل الإجراءات" if action_type == "all" else action_type

    # ✅ حساب عدد التقارير المتوقعة بعد الفلترة
    with SessionLocal() as session:
        count_query = session.query(Report).join(
            Patient, Report.patient_id == Patient.id
        ).filter(Patient.full_name.ilike(f"%{patient_name}%"))

        if action_type != "all":
            count_query = count_query.filter(Report.medical_action == action_type)

        preview_count = count_query.count()

    # عرض خيارات التاريخ مع ملخص
    keyboard = InlineKeyboardMarkup([
        [InlineKeyboardButton("📅 إضافة فلترة حسب التاريخ", callback_data="add_date_filter:yes")],
        [InlineKeyboardButton(f"✅ طباعة مباشرة ({preview_count} تقرير)", callback_data="add_date_filter:no")],
        [InlineKeyboardButton("🔙 رجوع", callback_data="back:action_type")],
        [InlineKeyboardButton("❌ إلغاء", callback_data="abort")],
    ])

    await q.edit_message_text(
        f"✅ **ملخص الاختيار:**\n\n"
        f"👤 **المريض:** {patient_name}\n"
        f"📌 **نوع الإجراء:** {action_label}\n"
        f"📊 **التقارير المتاحة:** {preview_count} تقرير\n\n"
        f"🔽 **هل تريد فلترة حسب التاريخ؟**",
        reply_markup=keyboard,
        parse_mode="Markdown"
    )
    return SELECT_YEAR


async def handle_manual_entry(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """معالجة زر الكتابة اليدوية"""
    q = update.callback_query
    await q.answer()
    
    keyboard = InlineKeyboardMarkup([
        [InlineKeyboardButton("🔙 رجوع", callback_data="back:filter")],
        [InlineKeyboardButton("❌ إلغاء", callback_data="abort")]
    ])
    
    await q.edit_message_text(
        "🏥 **طباعة تقارير مستشفى محدد**\n\n"
        "📝 أرسل اسم المستشفى:\n\n"
        "💡 يمكنك كتابة جزء من الاسم للبحث",
        reply_markup=keyboard,
        parse_mode="Markdown"
    )
    return ENTER_NAME

async def handle_back_from_enter_name(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """معالجة زر الرجوع من إدخال الاسم"""
    q = update.callback_query
    await q.answer()
    
    # ✅ مسح البيانات المؤقتة عند الرجوع
    context.user_data.pop("mode", None)
    context.user_data.pop("patients_list", None)
    context.user_data.pop("print_patient_page", None)
    context.user_data.pop("filter_value", None)
    context.user_data.pop("action_type", None)
    
    await q.edit_message_text("🖨️ اختر نوع الفلترة:", reply_markup=_filters_kb())
    return SELECT_FILTER

async def handle_department_selection(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """معالجة اختيار القسم من القائمة"""
    q = update.callback_query
    await q.answer()
    
    dept_name = q.data.split(":", 1)[1]
    context.user_data["department_value"] = dept_name
    
    hospital_name = context.user_data.get("filter_value")
    
    # الانتقال مباشرة لخيارات التاريخ
    keyboard = InlineKeyboardMarkup([
        [InlineKeyboardButton("📅 إضافة فلترة حسب التاريخ", callback_data="add_date_filter:yes")],
        [InlineKeyboardButton("✅ طباعة مباشرة (كل التواريخ)", callback_data="add_date_filter:no")],
        [InlineKeyboardButton("🔙 رجوع", callback_data="back:dept_list")],
        [InlineKeyboardButton("❌ إلغاء", callback_data="abort")],
    ])
    
    await q.edit_message_text(
        f"🏥 **المستشفى:** {hospital_name}\n"
        f"🏷️ **القسم:** {dept_name}\n\n"
        f"📋 هل تريد إضافة فلترة حسب التاريخ أيضاً؟",
        reply_markup=keyboard,
        parse_mode="Markdown"
    )
    
    return SELECT_YEAR

async def handle_dept_manual_entry(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """معالجة زر الكتابة اليدوية للقسم"""
    q = update.callback_query
    await q.answer()
    
    hospital_name = context.user_data.get("filter_value")
    
    keyboard = InlineKeyboardMarkup([
        [InlineKeyboardButton("🔙 رجوع", callback_data="back:dept_option")],
        [InlineKeyboardButton("❌ إلغاء", callback_data="abort")]
    ])
    
    await q.edit_message_text(
        f"🏥 **المستشفى:** {hospital_name}\n\n"
        f"🏷️ **أرسل اسم القسم:**\n\n"
        f"💡 يمكنك كتابة جزء من اسم القسم",
        reply_markup=keyboard,
        parse_mode="Markdown"
    )
    
    return ENTER_DEPARTMENT

async def handle_back_from_enter_dept(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """معالجة زر الرجوع من إدخال اسم القسم"""
    q = update.callback_query
    await q.answer()
    
    hospital_name = context.user_data.get("filter_value")
    
    keyboard = InlineKeyboardMarkup([
        [InlineKeyboardButton("🏥 المستشفى كاملاً", callback_data="dept_option:all")],
        [InlineKeyboardButton("🏷️ قسم معين فقط", callback_data="dept_option:specific")],
        [InlineKeyboardButton("🔙 رجوع", callback_data="back:filter")],
        [InlineKeyboardButton("❌ إلغاء", callback_data="abort")],
    ])
    
    await q.edit_message_text(
        f"🏥 **تم اختيار المستشفى:** {hospital_name}\n\n"
        f"📋 هل تريد طباعة:\n\n"
        f"🏥 **المستشفى كاملاً** (كل الأقسام)\n"
        f"🏷️ **قسم معين فقط** من المستشفى",
        reply_markup=keyboard,
        parse_mode="Markdown"
    )
    
    return SELECT_DEPARTMENT_OPTION

def export_to_pdf_windows(reports_data, filename="reports"):
    """تصدير التقارير إلى PDF احترافي مع دعم عربي كامل"""
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, KeepTogether
    from reportlab.lib.units import cm, mm
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from datetime import datetime
    from arabic_reshaper import reshape
    from bidi.algorithm import get_display
    import os
    
    try:
        # ✅ تسجيل خط عربي - جرب عدة خطوط من Windows و Linux
        arabic_font = 'Helvetica'
        import sys
        try:
            # الخيارات المتاحة في Windows (بالترتيب من الأفضل للأسوأ)
            font_options = []
            if sys.platform == 'win32':
                font_options = [
                    ("C:\\Windows\\Fonts\\tahoma.ttf", "Tahoma"),
                    ("C:\\Windows\\Fonts\\tahomabd.ttf", "TahomaBold"),
                    ("C:\\Windows\\Fonts\\arial.ttf", "Arial"),
                    ("C:\\Windows\\Fonts\\times.ttf", "Times"),
                ]
            else:
                # Linux font paths
                font_options = [
                    ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", "DejaVuSans"),
                    ("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", "DejaVuSansBold"),
                    ("/usr/share/fonts/truetype/freefont/FreeSans.ttf", "FreeSans"),
                    ("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf", "LiberationSans"),
                ]

            for font_path, font_name in font_options:
                if os.path.exists(font_path):
                    pdfmetrics.registerFont(TTFont(font_name, font_path))
                    arabic_font = font_name
                    logger.info(f"✅ تم تسجيل الخط العربي: {font_name}")
                    break
        except Exception as e:
            logger.warning(f"⚠️ فشل تسجيل الخط العربي، استخدام Helvetica: {e}")
            arabic_font = 'Helvetica'
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filepath = f"exports/{filename}_{timestamp}.pdf"
        
        # استخدام landscape للجدول الواسع
        doc = SimpleDocTemplate(
            filepath, 
            pagesize=landscape(A4),
            rightMargin=15*mm,
            leftMargin=15*mm,
            topMargin=20*mm,
            bottomMargin=20*mm
        )
        
        elements = []
        
        def ar(text):
            """معالجة النص العربي - RTL + إعادة تشكيل"""
            if not text:
                return ""
            try:
                clean_text = str(text).strip()
                reshaped = reshape(clean_text)
                bidi_text = get_display(reshaped)
                return bidi_text
            except Exception as e:
                return str(text)
        
        # ═══════════════════════════════════════
        # 🎨 صفحة الغلاف
        # ═══════════════════════════════════════
        
        # البسملة
        bismillah_style = ParagraphStyle(
            'Bismillah',
            fontSize=20,
            textColor=colors.HexColor('#27ae60'),
            alignment=TA_CENTER,
            spaceAfter=20*mm,
            fontName=arabic_font
        )
        elements.append(Paragraph(ar('بسم الله الرحمن الرحيم'), bismillah_style))
        
        # العنوان الرئيسي
        title_style = ParagraphStyle(
            'MainTitle',
            fontSize=28,
            textColor=colors.HexColor('#2c3e50'),
            alignment=TA_CENTER,
            spaceAfter=10*mm,
            fontName=arabic_font
        )
        elements.append(Paragraph(ar('التقرير الطبي الشامل'), title_style))
        
        # العنوان الفرعي
        subtitle_style = ParagraphStyle(
            'Subtitle',
            fontSize=16,
            textColor=colors.HexColor('#3498db'),
            alignment=TA_CENTER,
            spaceAfter=15*mm,
            fontName=arabic_font
        )
        elements.append(Paragraph(ar('نظام التقارير الطبية الذكي'), subtitle_style))
        
        # معلومات التقرير - جدول جميل
        # ✅ ترتيب RTL - القيمة ثم العنوان
        info_data = [
            [datetime.now().strftime("%Y-%m-%d %H:%M"), ar('📅 تاريخ الإنشاء')],
            [str(len(reports_data)), ar('📊 عدد التقارير')],
            [ar('نظام التقارير الطبية المتكامل'), ar('🏥 النظام')]
        ]

        info_table = Table(info_data, colWidths=[12*cm, 8*cm])
        info_table.setStyle(TableStyle([
            # ✅ RTL: العنوان يمين (العمود الأخير) والقيمة يسار
            ('BACKGROUND', (1, 0), (1, -1), colors.HexColor('#3498db')),
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#ecf0f1')),
            ('TEXTCOLOR', (1, 0), (1, -1), colors.white),
            ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#2c3e50')),
            ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),
            ('FONTNAME', (0, 0), (-1, -1), arabic_font),
            ('FONTSIZE', (0, 0), (-1, -1), 12),
            ('TOPPADDING', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 1.5, colors.HexColor('#2980b9')),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        
        elements.append(info_table)
        
        # ═══════════════════════════════════════
        # 📈 ملخص الإحصائيات الشامل
        # ═══════════════════════════════════════
        elements.append(Spacer(1, 10*mm))
        
        summary_title_style = ParagraphStyle(
            'SummaryTitle',
            fontSize=18,
            textColor=colors.HexColor('#2c3e50'),
            alignment=TA_RIGHT,
            spaceAfter=5*mm,
            fontName=arabic_font
        )
        elements.append(Paragraph(ar('📊 ملخص الإحصائيات الشامل'), summary_title_style))
        
        # حساب الإحصائيات
        stats = _calculate_detailed_stats(reports_data)
        
        # 1. المربعات الإحصائية العلوية
        summary_info_data = [
            [str(stats['unique_hospitals']), ar('🏥 عدد المستشفيات'), str(stats['unique_patients']), ar('🩺 عدد المرضى')],
            [str(stats.get('operations', 0)), ar('🔪 العمليات'), str(stats['total_reports']), ar('📄 إجمالي التقارير')]
        ]
        
        summary_info_table = Table(summary_info_data, colWidths=[4*cm, 6*cm, 4*cm, 6*cm])
        summary_info_table.setStyle(TableStyle([
            ('BACKGROUND', (1, 0), (1, -1), colors.HexColor('#3498db')),
            ('BACKGROUND', (3, 0), (3, -1), colors.HexColor('#3498db')),
            ('TEXTCOLOR', (1, 0), (1, -1), colors.white),
            ('TEXTCOLOR', (3, 0), (3, -1), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, -1), arabic_font),
            ('FONTSIZE', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#bdc3c7')),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        elements.append(summary_info_table)
        elements.append(Spacer(1, 10*mm))
        
        # 2. جدول توزيع الإجراءات
        by_action = stats.get('by_action', {})
        if by_action:
            elements.append(Paragraph(ar('توزيع الإجراءات الطبية:'), ParagraphStyle('ActionTitle', fontSize=14, fontName=arabic_font, alignment=TA_RIGHT, spaceAfter=3*mm)))
            
            summary_data = [[ar('النسبة'), ar('المجموع'), ar('نوع الإجراء')]]
            total = stats['total_reports']
            
            for action, count in sorted(by_action.items(), key=lambda x: x[1], reverse=True):
                percentage = f"{(count/total*100):.1f}%" if total > 0 else "0%"
                summary_data.append([percentage, str(count), ar(action)])
            
            summary_table = Table(summary_data, colWidths=[3*cm, 3*cm, 10*cm])
            summary_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, -1), arabic_font),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f9f9f9')]),
            ]))
            elements.append(summary_table)
        
        elements.append(PageBreak())
        
        # ═══════════════════════════════════════
        # 📋 جدول التقارير التفصيلي
        # ═══════════════════════════════════════
        
        section_title = ParagraphStyle(
            'SectionTitle',
            fontSize=20,
            textColor=colors.HexColor('#2c3e50'),
            alignment=TA_RIGHT,
            spaceAfter=5*mm,
            fontName=arabic_font,
            borderColor=colors.HexColor('#3498db'),
            borderWidth=2,
            borderPadding=5,
            backColor=colors.HexColor('#ecf0f1')
        )
        elements.append(Paragraph(ar('جدول التقارير التفصيلي'), section_title))
        elements.append(Spacer(1, 5*mm))
        
        # هيدر الجدول - ✅ ترتيب RTL (من اليمين لليسار)
        table_data = [[
            ar('الشكوى'),
            ar('الإجراء'),
            ar('الطبيب'),
            ar('القسم'),
            ar('المستشفى'),
            ar('المريض'),
            ar('التاريخ'),
            ar('م')
        ]]

        # البيانات - ✅ ترتيب RTL
        for idx, r in enumerate(reports_data, 1):
            complaint = r.get('complaint_text', '')
            if len(complaint) > 60:
                complaint = complaint[:57] + '...'

            row = [
                ar(complaint),
                ar(r.get('medical_action', 'غير محدد')[:20]),
                ar(r.get('doctor_name', 'غير محدد')[:20]),
                ar(r.get('department_name', 'غير محدد')[:20]),
                ar(r.get('hospital_name', 'غير محدد')[:25]),
                ar(r.get('patient_name', 'غير محدد')[:25]),
                r.get('report_date', '')[:10] if r.get('report_date') else '',
                str(idx)
            ]
            table_data.append(row)

        # إنشاء الجدول - ✅ ترتيب RTL
        reports_table = Table(table_data, colWidths=[
            5*cm,    # الشكوى
            3*cm,    # الإجراء
            3*cm,    # الطبيب
            3*cm,    # القسم
            4*cm,    # المستشفى
            3.5*cm,  # المريض
            2.5*cm,  # التاريخ
            1.5*cm   # م
        ])
        
        # تنسيق احترافي للجدول مع خط عربي
        reports_table.setStyle(TableStyle([
            # الهيدر
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), arabic_font),
            ('FONTSIZE', (0, 0), (-1, 0), 11),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),
            ('TOPPADDING', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            
            # البيانات - ✅ ترتيب RTL
            ('FONTNAME', (0, 1), (-1, -1), arabic_font),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
            ('ALIGN', (-1, 1), (-1, -1), 'CENTER'),  # م (الأخير = يمين في RTL)
            ('ALIGN', (-2, 1), (-2, -1), 'CENTER'),  # التاريخ
            ('ALIGN', (0, 1), (-3, -1), 'RIGHT'),    # باقي الأعمدة من اليمين
            ('VALIGN', (0, 1), (-1, -1), 'TOP'),
            ('TOPPADDING', (0, 1), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
            ('LEFTPADDING', (0, 1), (-1, -1), 5),
            ('RIGHTPADDING', (0, 1), (-1, -1), 5),
            
            # الحدود
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#bdc3c7')),
            ('LINEBELOW', (0, 0), (-1, 0), 2, colors.HexColor('#2980b9')),
            
            # تلوين الصفوف بالتناوب
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [
                colors.white,
                colors.HexColor('#f8f9fa')
            ]),
        ]))
        
        elements.append(reports_table)
        
        # إذا كان عدد التقارير كبير جداً
        if len(reports_data) > 100:
            elements.append(Spacer(1, 1*cm))
            note_style = ParagraphStyle(
                'Note', 
                fontSize=10, 
                textColor=colors.HexColor('#7f8c8d'), 
                alignment=TA_CENTER,
                fontName=arabic_font
            )
            elements.append(Paragraph(
                ar(f'ملاحظة: تم عرض جميع الـ {len(reports_data)} تقريراً في الجدول أعلاه'),
                note_style
            ))
        
        # بناء PDF
        doc.build(elements)
        
        logger.info(f"✅ تم إنشاء PDF احترافي: {filepath}")
        return filepath
        
    except Exception as e:
        logger.error(f"❌ خطأ في تصدير PDF: {e}")
        import traceback
        traceback.print_exc()
        return None

def _calculate_detailed_stats(reports_data):
    """حساب الإحصائيات التفصيلية من البيانات"""
    from collections import Counter
    
    stats = {
        'total_reports': len(reports_data),
        'unique_patients': len(set(r.get('patient_name') for r in reports_data if r.get('patient_name'))),
        'new_cases': 0,
        'operations': 0,
        'followups': 0,
        'medical_actions': 0,
        'unique_hospitals': len(set(r.get('hospital_name') for r in reports_data if r.get('hospital_name'))),
        'by_hospital': {},
        'by_department': {},
        'by_action': {},
        'by_doctor': {}
    }
    
    # عد حسب نوع الإجراء
    action_counter = Counter()
    hospital_counter = Counter()
    department_counter = Counter()
    doctor_counter = Counter()
    
    for r in reports_data:
        # دعم مفاتيح متعددة لضمان الحصول على البيانات
        action = r.get('medical_action') or r.get('action') or 'غير محدد'
        action_counter[action] += 1
        
        # تصنيف حسب نوع الإجراء
        if any(word in action for word in ['استشارة', 'جديد']):
            stats['new_cases'] += 1
        elif any(word in action for word in ['عملية', 'جراح']):
            stats['operations'] += 1
        elif any(word in action for word in ['متابعة', 'مراجعة']):
            stats['followups'] += 1
        elif any(word in action for word in ['إجراء', 'طبي']):
            stats['medical_actions'] += 1
        
        # عد حسب المستشفى
        hospital = r.get('hospital_name') or r.get('hospital')
        if hospital:
            hospital_counter[hospital] += 1
        
        # عد حسب القسم
        department = r.get('department_name') or r.get('department')
        if department:
            department_counter[department] += 1
        
        # عد حسب الطبيب
        doctor = r.get('doctor_name') or r.get('doctor') or r.get('doctor_full_name')
        if doctor:
            doctor_counter[doctor] += 1
    
    stats['by_hospital'] = dict(hospital_counter.most_common(10))
    stats['by_department'] = dict(department_counter.most_common(10))
    stats['by_action'] = dict(action_counter.most_common(10))
    stats['by_doctor'] = dict(doctor_counter.most_common(10))
    
    return stats

def _generate_dynamic_charts(reports_data, filter_type, stats):
    """إنشاء رسوم بيانية ديناميكية تتكيف حسب نوع الفلترة"""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    from matplotlib import font_manager
    import base64
    from io import BytesIO
    from arabic_reshaper import reshape
    from bidi.algorithm import get_display
    import os
    import sys

    # ✅ إعداد الخط العربي لـ matplotlib
    try:
        if sys.platform == 'win32':
            font_path = "C:\\Windows\\Fonts\\tahoma.ttf"
        else:
            # Linux font paths
            font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
            if not os.path.exists(font_path):
                font_path = "/usr/share/fonts/truetype/freefont/FreeSans.ttf"

        if os.path.exists(font_path):
            prop = font_manager.FontProperties(fname=font_path)
            plt.rcParams['font.family'] = prop.get_name()
        else:
            plt.rcParams['font.family'] = 'DejaVu Sans'
    except:
        plt.rcParams['font.family'] = 'DejaVu Sans'

    plt.rcParams['axes.unicode_minus'] = False
    
    def ar(text):
        """معالجة النص العربي لـ matplotlib"""
        try:
            reshaped = reshape(str(text))
            return get_display(reshaped)
        except:
            return str(text)
    
    def fig_to_base64(fig):
        """تحويل الرسم البياني إلى base64"""
        buffer = BytesIO()
        fig.savefig(buffer, format='png', dpi=150, bbox_inches='tight', facecolor='white')
        buffer.seek(0)
        image_base64 = base64.b64encode(buffer.read()).decode()
        plt.close(fig)
        return image_base64
    
    charts = {}
    
    # ═══════════════════════════════════════
    # 🎯 رسوم بيانية تتكيف حسب نوع الفلترة
    # ═══════════════════════════════════════
    
    logger.info(f"📊 إنشاء رسوم بيانية - نوع الفلتر: {filter_type}")
    
    # ═══════════════════════════════════════
    # 🎯 المنطق: اختيار الرسوم حسب نوع الفلترة
    # ═══════════════════════════════════════
    # - إذا الفلتر حسب المستشفى → رسوم تفصيلية للأقسام والإجراءات
    # - إذا الفلتر حسب القسم → رسوم تفصيلية للمستشفيات والإجراءات
    # - إذا الفلتر حسب الإجراء → رسوم تفصيلية للمستشفيات والأقسام
    # - إذا فلتر عام → جميع الرسوم
    
    # 1️⃣ رسم دائري للمستشفيات
    # يظهر دائماً إلا إذا كان الفلتر حسب مستشفى واحد فقط
    if stats['by_hospital'] and len(stats['by_hospital']) > 1:
        fig, ax = plt.subplots(figsize=(10, 7))
        hospitals = list(stats['by_hospital'].keys())[:8]
        values = list(stats['by_hospital'].values())[:8]
        
        colors = ['#3498db', '#e74c3c', '#2ecc71', '#f39c12', '#9b59b6', '#1abc9c', '#e67e22', '#34495e']
        
        wedges, texts, autotexts = ax.pie(
            values,
            labels=[ar(h) for h in hospitals],
            autopct='%1.1f%%',
            startangle=90,
            colors=colors,
            textprops={'fontsize': 12, 'weight': 'bold'}
        )
        
        for autotext in autotexts:
            autotext.set_color('white')
            autotext.set_fontsize(11)
        
        title = 'توزيع التقارير حسب المستشفيات'
        if filter_type == 'department':
            title += ' (للقسم المحدد)'
        elif filter_type and 'action' in str(filter_type).lower():
            title += ' (للإجراء المحدد)'
        
        ax.set_title(ar(title), fontsize=16, weight='bold', pad=20)
        charts['hospitals_pie'] = fig_to_base64(fig)
        logger.info("✅ تم إنشاء رسم المستشفيات")
    
    # 2️⃣ رسم أعمدة للأقسام
    # يظهر دائماً إلا إذا كان الفلتر حسب قسم واحد فقط
    if stats['by_department'] and len(stats['by_department']) > 0:
        # إذا كان قسم واحد فقط، لا نعرضه
        if len(stats['by_department']) > 1 or filter_type != 'department':
            fig, ax = plt.subplots(figsize=(12, 7))
            departments = list(stats['by_department'].keys())[:10]
            values = list(stats['by_department'].values())[:10]
            
            bars = ax.barh(range(len(departments)), values, color='#3498db', edgecolor='#2980b9', linewidth=1.5)
            
            # تدرج ألوان
            colors_gradient = ['#e74c3c', '#e67e22', '#f39c12', '#f1c40f', '#2ecc71', '#1abc9c', '#3498db', '#9b59b6', '#e84393', '#fd79a8']
            for i, bar in enumerate(bars):
                bar.set_color(colors_gradient[i % len(colors_gradient)])
            
            ax.set_yticks(range(len(departments)))
            ax.set_yticklabels([ar(d) for d in departments], fontsize=11)
            ax.set_xlabel(ar('عدد التقارير'), fontsize=13, weight='bold')
            
            title = 'توزيع التقارير حسب الأقسام'
            if filter_type == 'hospital':
                title += ' (في المستشفى المحدد)'
            elif filter_type and 'action' in str(filter_type).lower():
                title += ' (للإجراء المحدد)'
            
            ax.set_title(ar(title), fontsize=16, weight='bold', pad=20)
            ax.grid(axis='x', alpha=0.3, linestyle='--')
            
            # إضافة القيم على الأعمدة
            for i, v in enumerate(values):
                ax.text(v + 0.1, i, str(v), va='center', fontsize=10, weight='bold')
            
            charts['departments_bar'] = fig_to_base64(fig)
            logger.info("✅ تم إنشاء رسم الأقسام")
    
    # 3️⃣ رسم أعمدة للإجراءات الطبية
    # يظهر دائماً (مهم لكل أنواع الفلترة)
    if stats['by_action'] and len(stats['by_action']) > 0:
        fig, ax = plt.subplots(figsize=(12, 7))
        actions = list(stats['by_action'].keys())[:8]
        values = list(stats['by_action'].values())[:8]
        
        bars = ax.bar(range(len(actions)), values, color='#2ecc71', edgecolor='#27ae60', linewidth=1.5)
        
        # تدرج ألوان
        colors_gradient = ['#e74c3c', '#f39c12', '#2ecc71', '#3498db', '#9b59b6', '#1abc9c', '#e67e22', '#34495e']
        for i, bar in enumerate(bars):
            bar.set_color(colors_gradient[i % len(colors_gradient)])
        
        ax.set_xticks(range(len(actions)))
        ax.set_xticklabels([ar(a) for a in actions], fontsize=10, rotation=30, ha='right')
        ax.set_ylabel(ar('عدد التقارير'), fontsize=13, weight='bold')
        
        title = 'توزيع التقارير حسب الإجراءات الطبية'
        if filter_type == 'hospital':
            title += ' (في المستشفى المحدد)'
        elif filter_type == 'department':
            title += ' (في القسم المحدد)'
        
        ax.set_title(ar(title), fontsize=16, weight='bold', pad=20)
        ax.grid(axis='y', alpha=0.3, linestyle='--')
        
        # إضافة القيم على الأعمدة
        for i, v in enumerate(values):
            ax.text(i, v + 0.3, str(v), ha='center', fontsize=10, weight='bold')
        
        charts['actions_bar'] = fig_to_base64(fig)
        logger.info("✅ تم إنشاء رسم الإجراءات الطبية")
    
    # 4️⃣ رسم خط للتقارير حسب الوقت (إذا كانت البيانات تحتوي على تواريخ)
    try:
        from datetime import datetime
        from collections import defaultdict
        
        date_counter = defaultdict(int)
        for r in reports_data:
            date_str = r.get('report_date', '')
            if date_str and date_str != 'غير محدد':
                try:
                    # استخراج التاريخ فقط (بدون الوقت)
                    date_only = date_str.split()[0] if ' ' in date_str else date_str
                    date_counter[date_only] += 1
                except:
                    pass
        
        if date_counter and len(date_counter) > 1:
            fig, ax = plt.subplots(figsize=(14, 7))
            
            # ترتيب التواريخ
            sorted_dates = sorted(date_counter.items())
            dates = [d[0] for d in sorted_dates]
            counts = [d[1] for d in sorted_dates]
            
            ax.plot(range(len(dates)), counts, marker='o', linewidth=2.5, markersize=8, 
                   color='#3498db', markerfacecolor='#e74c3c', markeredgewidth=2, markeredgecolor='white')
            
            ax.fill_between(range(len(dates)), counts, alpha=0.3, color='#3498db')
            
            ax.set_xticks(range(len(dates)))
            ax.set_xticklabels(dates, fontsize=9, rotation=45, ha='right')
            ax.set_ylabel(ar('عدد التقارير'), fontsize=13, weight='bold')
            ax.set_title(ar('تطور عدد التقارير عبر الزمن'), fontsize=16, weight='bold', pad=20)
            ax.grid(True, alpha=0.3, linestyle='--')
            
            # إضافة القيم على النقاط
            for i, v in enumerate(counts):
                ax.text(i, v + 0.2, str(v), ha='center', fontsize=9, weight='bold')
            
            charts['timeline'] = fig_to_base64(fig)
    except Exception as e:
        logger.warning(f"⚠️ تعذر إنشاء رسم الزمن: {e}")
    
    # 5️⃣ رسم مقارنة الأقسام في كل مستشفى (Grouped Bar Chart)
    # يظهر فقط إذا كان هناك أكثر من مستشفى وأكثر من قسم
    try:
        from collections import defaultdict
        
        # تنظيم البيانات: {مستشفى: {قسم: عدد}}
        hospital_departments = defaultdict(lambda: defaultdict(int))
        for r in reports_data:
            hospital = r.get('hospital_name')
            department = r.get('department_name')
            if hospital and department:
                hospital_departments[hospital][department] += 1
        
        # عرض هذا الرسم فقط إذا لم يكن الفلتر حسب مستشفى واحد أو قسم واحد
        show_comparison = True
        if filter_type == 'hospital' and len(hospital_departments) == 1:
            show_comparison = False
        if filter_type == 'department':
            # إذا كان قسم واحد فقط، لا نعرض المقارنة
            all_depts = set()
            for _, depts in hospital_departments.items():
                all_depts.update(depts.keys())
            if len(all_depts) == 1:
                show_comparison = False
        
        if hospital_departments and len(hospital_departments) > 0 and show_comparison:
            # أخذ أفضل 4 مستشفيات
            top_hospitals = sorted(hospital_departments.items(), 
                                 key=lambda x: sum(x[1].values()), 
                                 reverse=True)[:4]
            
            # جمع كل الأقسام
            all_departments = set()
            for _, depts in top_hospitals:
                all_departments.update(depts.keys())
            
            if len(all_departments) > 0 and len(top_hospitals) > 1:
                fig, ax = plt.subplots(figsize=(14, 8))
                
                department_list = list(all_departments)[:8]  # أفضل 8 أقسام
                x = range(len(department_list))
                width = 0.2  # عرض كل عمود
                
                colors = ['#e74c3c', '#3498db', '#2ecc71', '#f39c12', '#9b59b6', '#1abc9c']
                
                for i, (hospital, depts) in enumerate(top_hospitals):
                    values = [depts.get(dept, 0) for dept in department_list]
                    offset = width * (i - len(top_hospitals)/2)
                    bars = ax.bar([pos + offset for pos in x], values, 
                                 width, label=ar(hospital),
                                 color=colors[i % len(colors)],
                                 edgecolor='white', linewidth=1.5)
                    
                    # إضافة القيم على الأعمدة
                    for j, (bar, val) in enumerate(zip(bars, values)):
                        if val > 0:
                            ax.text(bar.get_x() + bar.get_width()/2, val + 0.1, 
                                  str(val), ha='center', va='bottom', 
                                  fontsize=8, weight='bold')
                
                ax.set_xticks(x)
                ax.set_xticklabels([ar(d) for d in department_list], 
                                  fontsize=10, rotation=30, ha='right')
                ax.set_ylabel(ar('عدد التقارير'), fontsize=13, weight='bold')
                
                title = 'مقارنة الأقسام في المستشفيات'
                if filter_type == 'department':
                    title = f'توزيع {department_list[0] if len(department_list)==1 else "الأقسام"} حسب المستشفيات'
                
                ax.set_title(ar(title), fontsize=16, weight='bold', pad=20)
                ax.legend(fontsize=10, loc='upper right')
                ax.grid(axis='y', alpha=0.3, linestyle='--')
                
                charts['hospitals_departments_comparison'] = fig_to_base64(fig)
                logger.info("✅ تم إنشاء رسم المقارنة")
    except Exception as e:
        logger.warning(f"⚠️ تعذر إنشاء رسم المقارنة: {e}")
    
    # 6️⃣ خريطة حرارية (Heatmap) للمستشفيات والأقسام
    # مفيدة جداً للفلاتر العامة ولفلتر الإجراءات
    try:
        from collections import defaultdict
        import numpy as np
        
        # تنظيم البيانات
        hospital_departments = defaultdict(lambda: defaultdict(int))
        for r in reports_data:
            hospital = r.get('hospital_name')
            department = r.get('department_name')
            if hospital and department:
                hospital_departments[hospital][department] += 1
        
        # عرض الخريطة الحرارية إذا كان هناك تنوع
        show_heatmap = True
        if filter_type == 'hospital' and len(hospital_departments) == 1:
            show_heatmap = False
        
        if hospital_departments and len(hospital_departments) > 1 and show_heatmap:
            # أخذ أفضل المستشفيات والأقسام
            top_hospitals = sorted(hospital_departments.items(), 
                                 key=lambda x: sum(x[1].values()), 
                                 reverse=True)[:6]
            
            all_departments = set()
            for _, depts in top_hospitals:
                all_departments.update(depts.keys())
            department_list = sorted(list(all_departments))[:10]
            
            if len(department_list) > 0 and len(top_hospitals) > 1:
                # إنشاء مصفوفة البيانات
                data_matrix = []
                hospital_names = []
                
                for hospital, depts in top_hospitals:
                    row = [depts.get(dept, 0) for dept in department_list]
                    data_matrix.append(row)
                    hospital_names.append(hospital)
                
                data_matrix = np.array(data_matrix)
                
                fig, ax = plt.subplots(figsize=(14, 8))
                
                # رسم الخريطة الحرارية
                im = ax.imshow(data_matrix, cmap='YlOrRd', aspect='auto')
                
                # إعداد المحاور
                ax.set_xticks(range(len(department_list)))
                ax.set_yticks(range(len(hospital_names)))
                ax.set_xticklabels([ar(d) for d in department_list], 
                                  fontsize=10, rotation=45, ha='right')
                ax.set_yticklabels([ar(h) for h in hospital_names], fontsize=11)
                
                # إضافة القيم داخل الخلايا
                for i in range(len(hospital_names)):
                    for j in range(len(department_list)):
                        value = data_matrix[i, j]
                        if value > 0:
                            text_color = 'white' if value > data_matrix.max()/2 else 'black'
                            ax.text(j, i, str(int(value)), 
                                  ha='center', va='center',
                                  color=text_color, fontsize=10, weight='bold')
                
                title = 'خريطة توزيع التقارير (المستشفيات × الأقسام)'
                if filter_type == 'department':
                    title += ' (للقسم المحدد)'
                elif filter_type and 'action' in str(filter_type).lower():
                    title += ' (للإجراء المحدد)'
                
                ax.set_title(ar(title), fontsize=16, weight='bold', pad=20)
                
                # إضافة شريط الألوان
                cbar = plt.colorbar(im, ax=ax)
                cbar.set_label(ar('عدد التقارير'), rotation=0, labelpad=20)
                
                plt.tight_layout()
                charts['heatmap'] = fig_to_base64(fig)
                logger.info("✅ تم إنشاء الخريطة الحرارية")
    except Exception as e:
        logger.warning(f"⚠️ تعذر إنشاء الخريطة الحرارية: {e}")
    
    # 7️⃣ رسم نسب الأقسام في كل مستشفى (Stacked Bar)
    # يظهر للفلاتر العامة وفلتر القسم
    try:
        from collections import defaultdict
        
        hospital_departments = defaultdict(lambda: defaultdict(int))
        for r in reports_data:
            hospital = r.get('hospital_name')
            department = r.get('department_name')
            if hospital and department:
                hospital_departments[hospital][department] += 1
        
        # عرض هذا الرسم إذا كان هناك أكثر من مستشفى
        show_stacked = True
        if filter_type == 'hospital' and len(hospital_departments) == 1:
            show_stacked = False
        
        if hospital_departments and len(hospital_departments) > 1 and show_stacked:
            top_hospitals = sorted(hospital_departments.items(), 
                                 key=lambda x: sum(x[1].values()), 
                                 reverse=True)[:5]
            
            all_departments = set()
            for _, depts in top_hospitals:
                all_departments.update(depts.keys())
            department_list = sorted(list(all_departments))[:8]
            
            if len(department_list) > 0:
                fig, ax = plt.subplots(figsize=(12, 8))
                
                hospital_names = [h for h, _ in top_hospitals]
                colors = ['#e74c3c', '#3498db', '#2ecc71', '#f39c12', '#9b59b6', 
                         '#1abc9c', '#e67e22', '#34495e']
                
                # إنشاء الأعمدة المكدسة
                bottom = [0] * len(hospital_names)
                
                for i, dept in enumerate(department_list):
                    values = []
                    for _, depts in top_hospitals:
                        values.append(depts.get(dept, 0))
                    
                    bars = ax.barh(range(len(hospital_names)), values, 
                                  left=bottom, height=0.6,
                                  label=ar(dept), color=colors[i % len(colors)],
                                  edgecolor='white', linewidth=1)
                    
                    # تحديث القاع للطبقة التالية
                    bottom = [b + v for b, v in zip(bottom, values)]
                
                ax.set_yticks(range(len(hospital_names)))
                ax.set_yticklabels([ar(h) for h in hospital_names], fontsize=11)
                ax.set_xlabel(ar('عدد التقارير'), fontsize=13, weight='bold')
                
                title = 'نسب الأقسام في كل مستشفى'
                if filter_type == 'department':
                    title += ' (للقسم المحدد)'
                elif filter_type and 'action' in str(filter_type).lower():
                    title += ' (للإجراء المحدد)'
                
                ax.set_title(ar(title), fontsize=16, weight='bold', pad=20)
                ax.legend(fontsize=9, loc='center left', bbox_to_anchor=(1, 0.5))
                ax.grid(axis='x', alpha=0.3, linestyle='--')
                
                plt.tight_layout()
                charts['stacked_departments'] = fig_to_base64(fig)
                logger.info("✅ تم إنشاء رسم النسب المكدسة")
    except Exception as e:
        logger.warning(f"⚠️ تعذر إنشاء رسم النسب: {e}")
    
    return charts

def _get_filter_label(filter_type):
    """الحصول على تسمية نوع الفلترة"""
    labels = {
        'all': 'جميع التقارير',
        'today': 'تقارير اليوم',
        'week': 'تقارير هذا الأسبوع',
        'month': 'تقارير هذا الشهر',
        'custom': 'فترة مخصصة',
        'department': 'حسب القسم',
        'hospital': 'حسب المستشفى',
        'doctor': 'حسب الطبيب'
    }
    return labels.get(filter_type, '')

def export_to_html(reports_data, filename="reports", filter_type=None, title=None, filter_desc=None):
    """تصدير التقارير إلى HTML احترافي مع رسوم بيانية ديناميكية"""
    from jinja2 import Environment, FileSystemLoader
    from datetime import datetime
    import matplotlib
    matplotlib.use('Agg')  # للعمل بدون واجهة رسومية
    import matplotlib.pyplot as plt
    from collections import Counter
    import base64
    from io import BytesIO
    
    try:
        # تحضير البيانات للجدول - ندعم كلا صيغتي الصفوف (dicts من _query_reports أو dicts من مصادر أخرى)
        html_reports = []
        for r in reports_data:
            # دعم مفاتيح متعددة لضمان عدم فقدان البيانات
            report_id = r.get('id', r.get('report_id', ''))
            report_date = r.get('report_date') or r.get('date') or 'غير محدد'
            patient_name = r.get('patient_name') or r.get('patient') or r.get('patient_full_name') or 'غير محدد'
            hospital = r.get('hospital') or r.get('hospital_name') or 'غير محدد'
            department = r.get('department') or r.get('department_name') or 'غير محدد'
            doctor = r.get('doctor') or r.get('doctor_name') or r.get('doctor_full_name') or 'غير محدد'
            action = r.get('action') or r.get('medical_action') or 'غير محدد'
            complaint = r.get('complaint') or r.get('complaint_text') or 'لا يوجد'
            decision = r.get('decision') or r.get('doctor_decision') or 'لا يوجد'
            followup_date = r.get('followup_date') or ''
            followup_reason = r.get('followup_reason') or ''

            html_reports.append({
                'report_id': report_id,
                'report_date': report_date,
                'patient_name': patient_name,
                'hospital': hospital,
                'department': department,
                'doctor': doctor,
                'action': action,
                'complaint': complaint,
                'decision': decision,
                'followup_date': followup_date,
                'followup_reason': followup_reason
            })
        
        # ═══════════════════════════════════════
        # 📊 إنشاء الإحصائيات
        # ═══════════════════════════════════════
        
        stats = _calculate_detailed_stats(reports_data)
        
        # ═══════════════════════════════════════
        # 📈 إنشاء الرسوم البيانية الديناميكية
        # ═══════════════════════════════════════
        
        charts = _generate_dynamic_charts(reports_data, filter_type, stats)
        
        # ═══════════════════════════════════════
        # 🎨 تحميل القالب وإنشاء HTML
        # ═══════════════════════════════════════
        
        from jinja2 import TemplateNotFound
        env = Environment(loader=FileSystemLoader('templates'))
        try:
            template = env.get_template('report_summary.html')
        except TemplateNotFound:
            logger.warning("⚠️ قالب report_summary.html غير موجود في templates/ — سيتم استخدام قالب HTML افتراضي مبسّط")
            template = None
        
        context = {
            'title': title or 'التقرير الطبي الشامل',
            'generated_at': datetime.now().strftime("%Y-%m-%d %H:%M"),
            'filter_value': filter_desc or _get_filter_label(filter_type),
            'reports': html_reports,
            'stats': stats,
            'charts': charts
        }
        
        if template:
            html_content = template.render(**context)
        else:
            # قالب بديل بسيط في حال عدم وجود القالب الكامل
            rows_html = []
            for r in html_reports:
                rows_html.append(
                    f"<tr>"
                    f"<td>{r['report_id']}</td>"
                    f"<td>{r['report_date']}</td>"
                    f"<td>{r['patient_name']}</td>"
                    f"<td>{r['hospital']}</td>"
                    f"<td>{r['department']}</td>"
                    f"<td>{r['doctor']}</td>"
                    f"<td>{r['action']}</td>"
                    f"</tr>"
                )
            html_content = (
                f"<html lang=\"ar\" dir=\"rtl\"><head><meta charset=\"utf-8\"><title>{context['title']}</title>"
                f"<style>"
                f"* {{ direction: rtl; unicode-bidi: embed; }}"
                f"body {{ font-family: 'Segoe UI', Tahoma, Arial, sans-serif; direction: rtl; text-align: right; padding: 20px; color: #2c3e50; }}"
                f"h2 {{ text-align: center; color: #2c3e50; margin-bottom: 20px; }}"
                f"p {{ text-align: right; }}"
                f"table {{ width: 100%; border-collapse: collapse; direction: rtl; }}"
                f"th {{ background: #3498db; color: #fff; padding: 10px; text-align: right; }}"
                f"td {{ padding: 8px; text-align: right; border-bottom: 1px solid #ddd; }}"
                f"tr:nth-child(even) {{ background: #f8f9fa; }}"
                f"</style>"
                f"</head>"
                f"<body><h2>{context['title']}</h2><p>توليد: {context['generated_at']}</p>"
                f"<table>"
                f"<thead><tr><th>رقم</th><th>التاريخ</th><th>المريض</th><th>المستشفى</th><th>القسم</th><th>الطبيب</th><th>الإجراء</th></tr></thead>"
                f"<tbody>{''.join(rows_html)}</tbody></table></body></html>"
            )
        
        # حفظ الملف
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filepath = f"exports/{filename}_{timestamp}.html"
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        logger.info(f"✅ تم إنشاء HTML مع رسوم بيانية: {filepath}")
        return filepath
        
    except Exception as e:
        logger.error(f"❌ خطأ في تصدير HTML: {e}")
        import traceback
        traceback.print_exc()
        return None

def export_to_excel(reports_data, filename="reports"):
    """تصدير التقارير إلى Excel مع تنسيق احترافي"""
    import pandas as pd
    from datetime import datetime
    from openpyxl import load_workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    
    try:
        # تحضير البيانات - استخدام المفاتيح الصحيحة من _query_reports
        excel_data = []
        for r in reports_data:
            excel_data.append({
                'رقم': r.get('id', r.get('report_id', '')),
                'التاريخ': r.get('report_date', ''),
                'المريض': r.get('patient_name', ''),
                'المستشفى': r.get('hospital', r.get('hospital_name', '')),
                'القسم': r.get('department', r.get('department_name', '')),
                'الطبيب': r.get('doctor', r.get('doctor_name', '')),
                'الإجراء': r.get('action', r.get('medical_action', '')),
                'الشكوى': r.get('complaint', r.get('complaint_text', '')),
                'قرار الطبيب': r.get('decision', r.get('doctor_decision', '')),
                'موعد المراجعة': r.get('followup_date', ''),
                'سبب المراجعة': r.get('followup_reason', '')
            })
        
        # إنشاء DataFrame
        df = pd.DataFrame(excel_data)
        
        # حفظ مؤقتاً
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filepath = f"exports/{filename}_{timestamp}.xlsx"
        
        df.to_excel(filepath, index=False, engine='openpyxl', sheet_name='التقارير')
        
        # فتح الملف للتنسيق
        wb = load_workbook(filepath)
        ws = wb.active
        
        # حساب الإحصائيات للإضافة في البداية
        stats = _calculate_detailed_stats(reports_data)
        
        # إدراج صفوف للملخص في البداية (10 صفوف)
        ws.insert_rows(1, 10)
        
        # عنوان الملخص
        ws['A1'] = "📊 ملخص التقرير الإحصائي"
        ws['A1'].font = Font(bold=True, size=16, color="FFFFFF", name='Arial')
        ws['A1'].fill = PatternFill(start_color="2c3e50", end_color="2c3e50", fill_type="solid")
        ws['A1'].alignment = Alignment(horizontal="center", vertical="center")
        ws.merge_cells('A1:K1')
        
        # معلومات سريعة
        summary_labels = ["إجمالي التقارير", "عدد المرضى", "عدد المستشفيات", "العمليات"]
        summary_values = [stats['total_reports'], stats['unique_patients'], stats['unique_hospitals'], stats.get('operations', 0)]
        
        cols = ['B', 'D', 'G', 'I']
        for i, (label, value) in enumerate(zip(summary_labels, summary_values)):
            col = cols[i]
            # العنوان
            ws[f'{col}3'] = label
            ws[f'{col}3'].font = Font(bold=True, color="FFFFFF", size=11)
            ws[f'{col}3'].fill = PatternFill(start_color="3498db", end_color="3498db", fill_type="solid")
            ws[f'{col}3'].alignment = Alignment(horizontal="center")
            
            # القيمة
            ws[f'{col}4'] = value
            ws[f'{col}4'].font = Font(bold=True, size=14)
            ws[f'{col}4'].alignment = Alignment(horizontal="center")
            ws[f'{col}4'].border = Border(bottom=Side(style='medium', color='3498db'))

        # ملخص الإجراءات (أعلى 5)
        ws['A6'] = "🔝 أعلى 5 إجراءات طبية:"
        ws['A6'].font = Font(bold=True, size=12)
        ws.merge_cells('A6:C6')
        
        by_action = stats.get('by_action', {})
        for i, (action, count) in enumerate(list(sorted(by_action.items(), key=lambda x: x[1], reverse=True))[:5], 7):
            ws[f'A{i}'] = f"• {action}:"
            ws[f'B{i}'] = count
            ws[f'A{i}'].alignment = Alignment(horizontal="right")
            ws[f'B{i}'].font = Font(bold=True)
            
        # إضافة فاصل قبل الجدول
        ws['A10'] = "📋 تفاصيل التقارير:"
        ws['A10'].font = Font(bold=True, size=12, color="34495e")
        ws.merge_cells('A10:K10')
        
        # ✅ تفعيل RTL للورقة بالكامل
        ws.sheet_view.rightToLeft = True
        
        # تنسيق الهيدر (أصبح الآن في الصف 11 بعد الإدراج)
        header_row = 11
        header_fill = PatternFill(start_color="2c3e50", end_color="2c3e50", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=12, name='Arial')
        header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True, readingOrder=2)
        
        for cell in ws[header_row]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = header_alignment
        
        # تنسيق الخلايا - RTL
        thin_border = Border(
            left=Side(style='thin', color='bdc3c7'),
            right=Side(style='thin', color='bdc3c7'),
            top=Side(style='thin', color='bdc3c7'),
            bottom=Side(style='thin', color='bdc3c7')
        )
        
        for row in ws.iter_rows(min_row=header_row + 1, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
            for cell in row:
                cell.border = thin_border
                cell.alignment = Alignment(
                    horizontal="right", 
                    vertical="top", 
                    wrap_text=True,
                    readingOrder=2  # RTL
                )
                cell.font = Font(name='Arial', size=10)
        
        # تعديل عرض الأعمدة
        ws.column_dimensions['A'].width = 12  # رقم
        ws.column_dimensions['B'].width = 15  # التاريخ
        ws.column_dimensions['C'].width = 20  # المريض
        ws.column_dimensions['D'].width = 25  # المستشفى
        ws.column_dimensions['E'].width = 15  # القسم
        ws.column_dimensions['F'].width = 15  # الطبيب
        ws.column_dimensions['G'].width = 20  # الإجراء
        ws.column_dimensions['H'].width = 35  # الشكوى
        ws.column_dimensions['I'].width = 35  # القرار
        ws.column_dimensions['J'].width = 12  # موعد المراجعة
        ws.column_dimensions['K'].width = 20  # سبب المراجعة
        
        # تلوين الصفوف بالتناوب (يبدأ من بعد الهيدر)
        for idx, row in enumerate(ws.iter_rows(min_row=header_row + 1, max_row=ws.max_row), start=header_row + 1):
            if idx % 2 == 0:
                for cell in row:
                    cell.fill = PatternFill(start_color="f8f9fa", end_color="f8f9fa", fill_type="solid")
        
        # إضافة ورقة إحصائيات
        stats_ws = wb.create_sheet("الإحصائيات")
        stats_ws.sheet_view.rightToLeft = True  # RTL
        
        # عنوان
        stats_ws['A1'] = "📊 إحصائيات التقارير الشاملة"
        stats_ws['A1'].font = Font(bold=True, size=16, color="2c3e50", name='Arial')
        stats_ws['A1'].alignment = Alignment(horizontal="center", vertical="center")
        stats_ws['A1'].fill = PatternFill(start_color="ecf0f1", end_color="ecf0f1", fill_type="solid")
        stats_ws.merge_cells('A1:B1')
        
        # معلومات عامة
        labels = ["إجمالي التقارير:", "تاريخ الإنشاء:", "الحالة:"]
        values = [len(reports_data), datetime.now().strftime("%Y-%m-%d %H:%M"), "مكتمل ✅"]
        
        current_row = 3
        for label, value in zip(labels, values):
            stats_ws[f'A{current_row}'] = label
            stats_ws[f'B{current_row}'] = value
            stats_ws[f'A{current_row}'].font = Font(bold=True, size=11, name='Arial')
            stats_ws[f'A{current_row}'].alignment = Alignment(horizontal="right", vertical="center", readingOrder=2)
            stats_ws[f'B{current_row}'].alignment = Alignment(horizontal="right", vertical="center", readingOrder=2)
            stats_ws[f'B{current_row}'].font = Font(size=11, name='Arial')
            current_row += 1
            
        # ملخص الإجراءات الطبية
        current_row += 2
        stats_ws[f'A{current_row}'] = "📋 ملخص الإجراءات الطبية"
        stats_ws[f'A{current_row}'].font = Font(bold=True, size=14, color="FFFFFF", name='Arial')
        stats_ws[f'A{current_row}'].fill = PatternFill(start_color="3498db", end_color="3498db", fill_type="solid")
        stats_ws[f'A{current_row}'].alignment = Alignment(horizontal="center", vertical="center")
        stats_ws.merge_cells(f'A{current_row}:C{current_row}')
        current_row += 1
        
        stats_ws[f'A{current_row}'] = "نوع الإجراء"
        stats_ws[f'B{current_row}'] = "المجموع"
        stats_ws[f'C{current_row}'] = "النسبة المئوية"
        stats_ws[f'A{current_row}'].font = Font(bold=True, size=11, name='Arial')
        stats_ws[f'B{current_row}'].font = Font(bold=True, size=11, name='Arial')
        stats_ws[f'C{current_row}'].font = Font(bold=True, size=11, name='Arial')
        stats_ws[f'A{current_row}'].alignment = Alignment(horizontal="center", vertical="center")
        stats_ws[f'B{current_row}'].alignment = Alignment(horizontal="center", vertical="center")
        stats_ws[f'C{current_row}'].alignment = Alignment(horizontal="center", vertical="center")
        current_row += 1
        
        # حساب الإحصائيات
        stats = _calculate_detailed_stats(reports_data)
        by_action = stats.get('by_action', {})
        total_reports = len(reports_data)
        
        for action, count in sorted(by_action.items(), key=lambda x: x[1], reverse=True):
            percentage = (count / total_reports * 100) if total_reports > 0 else 0
            stats_ws[f'A{current_row}'] = action
            stats_ws[f'B{current_row}'] = count
            stats_ws[f'C{current_row}'] = f"{percentage:.1f}%"
            stats_ws[f'A{current_row}'].alignment = Alignment(horizontal="right", vertical="center", readingOrder=2)
            stats_ws[f'B{current_row}'].alignment = Alignment(horizontal="center", vertical="center")
            stats_ws[f'C{current_row}'].alignment = Alignment(horizontal="center", vertical="center")
            stats_ws[f'A{current_row}'].font = Font(name='Arial', size=10)
            stats_ws[f'B{current_row}'].font = Font(name='Arial', size=10)
            stats_ws[f'C{current_row}'].font = Font(name='Arial', size=10)
            current_row += 1
            
        # صف الإجمالي النهائي
        stats_ws[f'A{current_row}'] = "الإجمالي الكلي"
        stats_ws[f'B{current_row}'] = total_reports
        stats_ws[f'C{current_row}'] = "100%"
        stats_ws[f'A{current_row}'].font = Font(bold=True, size=11, color="FFFFFF", name='Arial')
        stats_ws[f'B{current_row}'].font = Font(bold=True, size=11, color="FFFFFF", name='Arial')
        stats_ws[f'C{current_row}'].font = Font(bold=True, size=11, color="FFFFFF", name='Arial')
        stats_ws[f'A{current_row}'].fill = PatternFill(start_color="2c3e50", end_color="2c3e50", fill_type="solid")
        stats_ws[f'B{current_row}'].fill = PatternFill(start_color="2c3e50", end_color="2c3e50", fill_type="solid")
        stats_ws[f'C{current_row}'].fill = PatternFill(start_color="2c3e50", end_color="2c3e50", fill_type="solid")
        stats_ws[f'A{current_row}'].alignment = Alignment(horizontal="right", vertical="center")
        stats_ws[f'B{current_row}'].alignment = Alignment(horizontal="center", vertical="center")
        stats_ws[f'C{current_row}'].alignment = Alignment(horizontal="center", vertical="center")

        stats_ws.column_dimensions['A'].width = 30
        stats_ws.column_dimensions['B'].width = 15
        stats_ws.column_dimensions['C'].width = 15
        
        wb.save(filepath)
        wb.close()
        
        logger.info(f"✅ تم إنشاء Excel: {filepath}")
        return filepath
        
    except Exception as e:
        logger.error(f"❌ خطأ في تصدير Excel: {e}")
        return None

def export_to_word(reports_data, filename="reports"):
    """تصدير التقارير إلى Word مع تنسيق احترافي"""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from datetime import datetime
    
    try:
        doc = Document()

        # ✅ إعداد RTL على مستوى المستند بالكامل
        from docx.oxml import OxmlElement
        # جعل كل الأقسام RTL
        for section in doc.sections:
            sectPr = section._sectPr
            bidi_elem = OxmlElement('w:bidi')
            bidi_elem.set(qn('w:val'), '1')
            sectPr.append(bidi_elem)

        # إعداد الخط العربي والـ RTL
        doc.styles['Normal'].font.name = 'Arial'
        doc.styles['Normal'].font.size = Pt(11)
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        doc.styles['Normal'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # ✅ تفعيل RTL على مستوى الفقرة الافتراضية
        pPr = doc.styles['Normal']._element.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)
        
        # صفحة الغلاف
        title = doc.add_heading('بسم الله الرحمن الرحيم', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.color.rgb = RGBColor(39, 174, 96)
        
        main_title = doc.add_heading('📊 التقرير الطبي الشامل', 0)
        main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        main_title.runs[0].font.color.rgb = RGBColor(52, 152, 219)
        
        subtitle = doc.add_paragraph('نظام التقارير الطبية الذكي')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(14)
        subtitle.runs[0].font.color.rgb = RGBColor(127, 140, 141)
        
        doc.add_paragraph()  # مسافة
        
        # معلومات التقرير
        info_table = doc.add_table(rows=2, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        info_table.cell(0, 0).text = '📅 تاريخ الإنشاء'
        info_table.cell(0, 1).text = datetime.now().strftime("%Y-%m-%d %H:%M")
        info_table.cell(1, 0).text = '📊 عدد التقارير'
        info_table.cell(1, 1).text = str(len(reports_data))
        
        for row in info_table.rows:
            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()
        
        # 📈 ملخص الإجراءات الطبية (جديد)
        stats_title = doc.add_heading('📈 ملخص الإجراءات الطبية', level=1)
        stats_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        stats_title.runs[0].font.color.rgb = RGBColor(52, 152, 219)
        
        # حساب الإحصائيات
        stats = _calculate_detailed_stats(reports_data)
        by_action = stats.get('by_action', {})
        
        if by_action:
            action_table = doc.add_table(rows=1, cols=2)
            action_table.style = 'Table Grid'
            
            # الهيدر
            hdr_cells = action_table.rows[0].cells
            hdr_cells[1].text = 'نوع الإجراء'
            hdr_cells[0].text = 'المجموع'
            
            for cell in hdr_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # تفعيل RTL للهيدر
                pPr = cell.paragraphs[0]._element.get_or_add_pPr()
                bidi = OxmlElement('w:bidi')
                bidi.set(qn('w:val'), '1')
                pPr.append(bidi)
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.bold = True
            
            # البيانات
            for action, count in sorted(by_action.items(), key=lambda x: x[1], reverse=True):
                row_cells = action_table.add_row().cells
                row_cells[1].text = action
                row_cells[0].text = str(count)
                
                for cell in row_cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    # تفعيل RTL للبيانات
                    pPr = cell.paragraphs[0]._element.get_or_add_pPr()
                    bidi = OxmlElement('w:bidi')
                    bidi.set(qn('w:val'), '1')
                    pPr.append(bidi)
                
            # صف الإجمالي
            total_row = action_table.add_row().cells
            total_row[1].text = 'إجمالي الإجراءات'
            total_row[0].text = str(len(reports_data))
            
            for cell in total_row:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                pPr = cell.paragraphs[0]._element.get_or_add_pPr()
                bidi = OxmlElement('w:bidi')
                bidi.set(qn('w:val'), '1')
                pPr.append(bidi)
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.bold = True

        doc.add_page_break()
        
        # التقارير التفصيلية
        for idx, report in enumerate(reports_data, 1):
            # عنوان التقرير
            heading = doc.add_heading(f'التقرير رقم {idx}', level=1)
            heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)
            
            # جدول المعلومات الأساسية
            table = doc.add_table(rows=6, cols=2)
            table.style = 'Medium Shading 1 Accent 1'
            
            table.cell(0, 0).text = '📅 التاريخ'
            table.cell(0, 1).text = report.get('report_date', 'غير محدد')
            
            table.cell(1, 0).text = '👤 المريض'
            table.cell(1, 1).text = report.get('patient_name', 'غير محدد')
            
            table.cell(2, 0).text = '🏥 المستشفى'
            table.cell(2, 1).text = report.get('hospital', report.get('hospital_name', 'غير محدد'))
            
            table.cell(3, 0).text = '🏢 القسم'
            table.cell(3, 1).text = report.get('department', report.get('department_name', 'غير محدد'))
            
            table.cell(4, 0).text = '👨‍⚕️ الطبيب'
            table.cell(4, 1).text = report.get('doctor', report.get('doctor_name', 'غير محدد'))
            
            table.cell(5, 0).text = '⚕️ الإجراء'
            table.cell(5, 1).text = report.get('action', report.get('medical_action', 'غير محدد'))
            
            # محاذاة النص لليمين - RTL
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        # ✅ تفعيل RTL على كل فقرة داخل الجدول
                        pPr = paragraph._element.get_or_add_pPr()
                        bidi = OxmlElement('w:bidi')
                        bidi.set(qn('w:val'), '1')
                        pPr.append(bidi)
                        for run in paragraph.runs:
                            run.font.name = 'Arial'
                            # ✅ تفعيل RTL على مستوى الـ run
                            rPr = run._element.get_or_add_rPr()
                            rtl = OxmlElement('w:rtl')
                            rtl.set(qn('w:val'), '1')
                            rPr.append(rtl)
                    # العمود الأول bold
                    if cell == row.cells[0]:
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(44, 62, 80)
            
            doc.add_paragraph()
            
            # دالة مساعدة لتفعيل RTL على فقرة
            def _set_paragraph_rtl(p):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                pPr = p._element.get_or_add_pPr()
                bidi = OxmlElement('w:bidi')
                bidi.set(qn('w:val'), '1')
                pPr.append(bidi)

            # الشكوى
            h_complaint = doc.add_heading('📝 الشكوى:', level=2)
            _set_paragraph_rtl(h_complaint)
            complaint_p = doc.add_paragraph(report.get('complaint', report.get('complaint_text', 'لا يوجد')))
            _set_paragraph_rtl(complaint_p)

            # قرار الطبيب
            h_decision = doc.add_heading('✅ قرار الطبيب:', level=2)
            _set_paragraph_rtl(h_decision)
            decision_p = doc.add_paragraph(report.get('decision', report.get('doctor_decision', 'لا يوجد')))
            _set_paragraph_rtl(decision_p)

            # موعد المراجعة
            if report.get('followup_date'):
                h_followup = doc.add_heading('📅 موعد المراجعة:', level=2)
                _set_paragraph_rtl(h_followup)
                followup_p = doc.add_paragraph()
                followup_p.add_run(f"التاريخ: {report.get('followup_date', '')}\n").bold = True
                followup_p.add_run(f"السبب: {report.get('followup_reason', '')}")
                _set_paragraph_rtl(followup_p)
            
            # فاصل
            if idx < len(reports_data):
                doc.add_page_break()
        
        # حفظ الملف
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filepath = f"exports/{filename}_{timestamp}.docx"
        
        doc.save(filepath)
        
        logger.info(f"✅ تم إنشاء Word: {filepath}")
        return filepath
        
    except Exception as e:
        logger.error(f"❌ خطأ في تصدير Word: {e}")
        return None

# ================================================
# معالجات طباعة حسب المريض (نص داخل Telegram)
# ================================================

async def show_patient_selection_for_print_local(query, context):
    """عرض قائمة المرضى لاختيار مريض للطباعة (نسخة محلية لـ admin_reports)"""
    try:
        from services.patients_service import get_all_patients
        from telegram.constants import ParseMode
        
        patients = get_all_patients()
        
        if not patients:
            await query.edit_message_text(
                "⚠️ **لا توجد مرضى في قاعدة البيانات**\n\n"
                "يرجى إضافة مرضى أولاً.",
                parse_mode=ParseMode.MARKDOWN
            )
            return ConversationHandler.END
        
        # ترتيب المرضى أبجدياً
        patients_sorted = sorted(patients, key=lambda x: x.get('name', ''))
        
        # تقسيم المرضى إلى صفحات (10 لكل صفحة)
        page = context.user_data.get('print_patient_page', 0)
        items_per_page = 10
        start_idx = page * items_per_page
        end_idx = start_idx + items_per_page
        patients_page = patients_sorted[start_idx:end_idx]
        
        keyboard_buttons = []
        
        # أزرار المرضى
        for patient in patients_page:
            patient_name = patient.get('name', 'غير محدد')
            patient_id = patient.get('id')
            keyboard_buttons.append([
                InlineKeyboardButton(
                    f"👤 {patient_name}",
                    callback_data=f"print_patient:{patient_id}"
                )
            ])
        
        # أزرار التنقل
        nav_buttons = []
        if page > 0:
            nav_buttons.append(InlineKeyboardButton("◀️ السابق", callback_data="patient_page:prev"))
        if end_idx < len(patients_sorted):
            nav_buttons.append(InlineKeyboardButton("التالي ▶️", callback_data="patient_page:next"))
        if nav_buttons:
            keyboard_buttons.append(nav_buttons)
        
        # أزرار الإلغاء والرجوع
        keyboard_buttons.append([
            InlineKeyboardButton("🔙 رجوع", callback_data="back:filter"),
            InlineKeyboardButton("❌ إلغاء", callback_data="abort")
        ])
        
        keyboard = InlineKeyboardMarkup(keyboard_buttons)
        
        text = f"""
🖨️ **طباعة حسب المريض**

اختر المريض المطلوب:
📊 إجمالي المرضى: {len(patients_sorted)}
📄 الصفحة: {page + 1} من {(len(patients_sorted) + items_per_page - 1) // items_per_page}
"""
        
        await query.edit_message_text(text, reply_markup=keyboard, parse_mode=ParseMode.MARKDOWN)
        return PRINT_SELECT_PATIENT
        
    except Exception as e:
        logger.error(f"❌ خطأ في عرض قائمة المرضى: {e}", exc_info=True)
        await query.edit_message_text(
            f"❌ **حدث خطأ**\n\n{str(e)}",
            parse_mode=ParseMode.MARKDOWN
        )
        return ConversationHandler.END


async def handle_patient_text_selection(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """معالجة اختيار المريض للطباعة كنص"""
    query = update.callback_query
    await query.answer()
    
    if query.data == "back:filter":
        await query.edit_message_text("🖨️ اختر نوع الفلترة:", reply_markup=_filters_kb())
        return SELECT_FILTER
    
    if query.data == "abort" or query.data == "print:cancel":
        await query.edit_message_text("❌ تم إلغاء الطباعة")
        return ConversationHandler.END
    
    if query.data.startswith("patient_page:"):
        # التنقل بين الصفحات
        direction = query.data.split(":")[1]
        current_page = context.user_data.get('print_patient_page', 0)
        if direction == "next":
            context.user_data['print_patient_page'] = current_page + 1
        elif direction == "prev":
            context.user_data['print_patient_page'] = max(0, current_page - 1)
        return await show_patient_selection_for_print_local(query, context)
    
    if query.data.startswith("print_patient:"):
        # اختيار المريض
        patient_id = int(query.data.split(":")[1])
        context.user_data['selected_patient_id'] = patient_id
        
        await query.edit_message_text("⏳ **جاري تحضير التقرير...**\n\nقد يستغرق هذا بضع ثوانٍ...")
        
        # طباعة تقارير المريض
        from bot.handlers.admin.admin_printing import print_patient_reports_as_text
        return await print_patient_reports_as_text(query, context, patient_id)
    
    return PRINT_SELECT_PATIENT


async def handle_patient_text_page(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """معالجة التنقل بين صفحات المرضى"""
    return await handle_patient_text_selection(update, context)

def register(app):
    conv = ConversationHandler(
       entry_points=[
    MessageHandler(filters.Regex("(طباعة التقارير|🖨️ طباعة التقارير)"), start_reports_filter)
           ],
        states={
            SELECT_FILTER: [
                CallbackQueryHandler(handle_filter_choice, pattern=r"^filter:(patient|patient_text|hospital|department|date|all)$"),
                CallbackQueryHandler(handle_print_patient_options, pattern=r"^print_patient_(all|period)$"),
                # معالجات لطباعة حسب المريض (نص داخل Telegram)
                CallbackQueryHandler(handle_patient_text_selection, pattern=r"^print_patient:\d+$"),
                CallbackQueryHandler(handle_patient_text_page, pattern=r"^patient_page:(next|prev)$"),
                CallbackQueryHandler(handle_back_to_filter, pattern=r"^back:type$"),
                CallbackQueryHandler(handle_back_to_filter, pattern=r"^back:filter$"),
                CallbackQueryHandler(confirm_export, pattern=r"^abort$"),
                CallbackQueryHandler(confirm_export, pattern=r"^print:cancel$"),
                MessageHandler(filters.Regex("^❌ إلغاء المحادثة$"), cancel_text),
            ],
            ENTER_NAME: [
                CallbackQueryHandler(handle_print_patient_selection, pattern=r"^print_patient:\d+$"),
                CallbackQueryHandler(handle_patient_page, pattern=r"^patient_page:\d+$"),
                CallbackQueryHandler(handle_hospital_selection, pattern=r"^select_hospital:"),
                CallbackQueryHandler(handle_department_selection, pattern=r"^select_dept:"),
                CallbackQueryHandler(handle_manual_entry, pattern=r"^(hospital|dept):manual$"),
                CallbackQueryHandler(handle_manual_entry, pattern=r"^hospital:manual$"),
                MessageHandler(filters.TEXT & ~filters.COMMAND, get_name),
                CallbackQueryHandler(handle_back_from_enter_name, pattern=r"^back:filter$"),
                CallbackQueryHandler(confirm_export, pattern=r"^abort$"),
                MessageHandler(filters.Regex("^❌ إلغاء المحادثة$"), cancel_text),
            ],
            SELECT_DEPARTMENT_OPTION: [
                CallbackQueryHandler(handle_department_option, pattern=r"^(dept_option:|back:filter)"),
                CallbackQueryHandler(confirm_export, pattern=r"^abort$"),
                MessageHandler(filters.Regex("^❌ إلغاء المحادثة$"), cancel_text),
            ],
            ENTER_DEPARTMENT: [
                CallbackQueryHandler(handle_department_selection, pattern=r"^select_dept:"),
                CallbackQueryHandler(handle_dept_manual_entry, pattern=r"^dept:manual$"),
                MessageHandler(filters.TEXT & ~filters.COMMAND, handle_department_name),
                CallbackQueryHandler(handle_back_from_enter_dept, pattern=r"^back:dept_option$"),
                CallbackQueryHandler(confirm_export, pattern=r"^abort$"),
                MessageHandler(filters.Regex("^❌ إلغاء المحادثة$"), cancel_text),
            ],
            SELECT_ACTION_TYPE: [
                CallbackQueryHandler(handle_action_type_choice, pattern=r"^action_type:"),
                CallbackQueryHandler(handle_action_type_choice, pattern=r"^back:patient_list$"),
                CallbackQueryHandler(confirm_export, pattern=r"^abort$"),
                MessageHandler(filters.Regex("^❌ إلغاء المحادثة$"), cancel_text),
            ],
            SELECT_YEAR: [
                CallbackQueryHandler(handle_year_choice, pattern=r"^(year:|add_date_filter:|back:filter|back:dept_option|back:dept_list|back:action_type)"),
                CallbackQueryHandler(confirm_export, pattern=r"^abort$"),
                MessageHandler(filters.Regex("^❌ إلغاء المحادثة$"), cancel_text),
            ],
            SELECT_MONTH: [
                CallbackQueryHandler(handle_month_choice, pattern=r"^(month:|back:year)"),
                CallbackQueryHandler(confirm_export, pattern=r"^abort$"),
                MessageHandler(filters.Regex("^❌ إلغاء المحادثة$"), cancel_text),
            ],
            CONFIRM_EXPORT: [
                CallbackQueryHandler(confirm_export, pattern=r"^(export:(yes|no|pdf|excel|word|html)|abort|back:confirm)$"),
                MessageHandler(filters.Regex("^❌ إلغاء المحادثة$"), cancel_text),
            ],
            PRINT_SELECT_PATIENT: [
                CallbackQueryHandler(handle_patient_text_selection, pattern=r"^print_patient:\d+$"),
                CallbackQueryHandler(handle_patient_text_page, pattern=r"^patient_page:(next|prev)$"),
                CallbackQueryHandler(handle_back_to_filter, pattern=r"^back:filter$"),
                CallbackQueryHandler(confirm_export, pattern=r"^print:cancel$"),
                CallbackQueryHandler(confirm_export, pattern=r"^abort$"),
                MessageHandler(filters.Regex("^❌ إلغاء المحادثة$"), cancel_text),
            ],
        },
        fallbacks=[
            CallbackQueryHandler(confirm_export, pattern=r"^abort$"),
            MessageHandler(filters.Regex("^❌ إلغاء المحادثة$"), cancel_text),
            MessageHandler(filters.Regex("^إلغاء$|^الغاء$|^cancel$"), cancel_text),
        ],
        name="admin_reports_conv",
        per_chat=True,
        per_user=True,
        per_message=False,
    )
    app.add_handler(conv)
    
    # ✅ تسجيل أمر /print_patient للبحث الفوري
    from telegram.ext import CommandHandler
    app.add_handler(CommandHandler("print_patient", handle_print_patient_command))
